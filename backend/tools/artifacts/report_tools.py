from __future__ import annotations

import json
import os
import subprocess
import sys
import time
from pathlib import Path
from typing import Any, Dict, List

from backend.tools.coding.foundation.core_mechanisms.path_security import (
    PathSecurityError,
    get_workspace_root,
    safe_path_for_read,
    safe_path_for_write,
)

from .docx_tools import docx_inspect_layout


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
TEXT_EXTENSIONS = {".txt", ".csv", ".json", ".md", ".log"}
MAX_CAPTURE_CHARS = 1600
MAX_CODE_CHARS = 8000


def office_report_from_code_run(
    output_path: str,
    title: str = "",
    assignment: str = "",
    code: str = "",
    script_path: str = "",
    command: str = "",
    working_dir: str = "",
    artifacts_dir: str = "",
    timeout: int = 120,
    render: bool = False,
    conclusion: str = "",
    language: str = "python",
) -> str:
    """Run code/command in the background and assemble a DOCX report with evidence."""
    started = time.time()
    activity: List[Dict[str, Any]] = []
    try:
        target = _safe_write_path(output_path, "output_path")
        work_dir = _working_dir(working_dir)
        artifact_dir = _artifact_dir(artifacts_dir, target)
    except PathSecurityError as exc:
        return _json_error("Path safety denied", activity=activity, detail=str(exc))

    _add_activity(activity, "prepare", "Prepare workspace", True, path=str(work_dir))
    before = _snapshot_files(artifact_dir)
    script = None
    source_code = ""
    command_display = ""
    run_result: Dict[str, Any]

    try:
        script, source_code, command_display = _prepare_execution(
            code=code,
            script_path=script_path,
            command=command,
            language=language,
            artifact_dir=artifact_dir,
            activity=activity,
        )
        run_result = _run_execution(
            script=script,
            command=command,
            work_dir=work_dir,
            artifact_dir=artifact_dir,
            output_path=target,
            timeout=max(1, int(timeout or 120)),
            activity=activity,
        )
    except Exception as exc:
        run_result = {
            "ok": False,
            "returncode": None,
            "stdout": "",
            "stderr": f"{type(exc).__name__}: {exc}",
            "duration_ms": 0,
            "timed_out": False,
        }
        _add_activity(activity, "run_code", "Run code", False, detail=run_result["stderr"])

    artifacts = _collect_artifacts(artifact_dir, before)
    _add_activity(
        activity,
        "collect_artifacts",
        "Collect generated artifacts",
        True,
        detail=f"{len(artifacts)} file(s)",
    )

    report_ok = False
    report_error = ""
    try:
        _write_docx_report(
            target=target,
            title=title or "Code Execution Report",
            assignment=assignment,
            source_code=source_code,
            command_display=command_display,
            run_result=run_result,
            artifacts=artifacts,
            conclusion=conclusion,
        )
        report_ok = True
        _add_activity(activity, "write_report", "Write DOCX report", True, path=str(target))
    except Exception as exc:
        report_error = f"{type(exc).__name__}: {exc}"
        _add_activity(activity, "write_report", "Write DOCX report", False, detail=report_error)

    inspect_payload: Dict[str, Any] = {}
    if report_ok:
        try:
            inspect_payload = json.loads(
                docx_inspect_layout(str(target), render=bool(render), output_dir=str(target.parent / f"{target.stem}_render"))
            )
        except Exception as exc:
            inspect_payload = {"ok": False, "error": f"{type(exc).__name__}: {exc}"}
        _add_activity(
            activity,
            "verify_report",
            "Inspect/render report",
            bool(inspect_payload.get("ok")),
            detail=_verification_detail(inspect_payload),
        )

    ok = bool(report_ok)
    status = "done"
    if not run_result.get("ok"):
        status = "code_failed_report_written" if report_ok else "code_failed"
    if report_ok and render and not _render_ok(inspect_payload):
        status = "render_unavailable_or_failed"

    payload = {
        "ok": ok,
        "status": status,
        "schema": "metis.artifact.code_report.v1",
        "output_path": str(target) if report_ok else "",
        "script_path": str(script) if script else "",
        "working_dir": str(work_dir),
        "artifacts_dir": str(artifact_dir),
        "command": command_display,
        "returncode": run_result.get("returncode"),
        "timed_out": bool(run_result.get("timed_out")),
        "duration_ms": int((time.time() - started) * 1000),
        "stdout": _truncate(run_result.get("stdout", "")),
        "stderr": _truncate(run_result.get("stderr", "")),
        "artifacts": artifacts,
        "report_inspection": inspect_payload,
        "report_error": report_error,
        "activity": activity,
        "artifact_activity": {
            "kind": "code_to_report",
            "summary": _activity_summary(status, target if report_ok else None, artifacts),
            "items": activity,
            "artifacts": artifacts,
            "output_path": str(target) if report_ok else "",
        },
    }
    return _json(payload)


def _prepare_execution(
    *,
    code: str,
    script_path: str,
    command: str,
    language: str,
    artifact_dir: Path,
    activity: List[Dict[str, Any]],
) -> tuple[Path | None, str, str]:
    inline_code = str(code or "")
    if inline_code.strip():
        if str(language or "python").strip().lower() not in {"python", "py"}:
            raise ValueError("Inline code execution only supports Python in v1.")
        script = artifact_dir / "analysis.py"
        script.write_text(inline_code, encoding="utf-8")
        _add_activity(activity, "write_code", "Write Python script", True, path=str(script))
        return script, inline_code, f"{sys.executable} {script}"

    if str(script_path or "").strip():
        script = safe_path_for_read(script_path)
        try:
            source_code = script.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            source_code = ""
        _add_activity(activity, "load_code", "Load script", True, path=str(script))
        return script, source_code, f"{sys.executable} {script}"

    if str(command or "").strip():
        _add_activity(activity, "load_command", "Use provided command", True, detail=str(command).strip())
        return None, "", str(command).strip()

    raise ValueError("Provide code, script_path, or command.")


def _run_execution(
    *,
    script: Path | None,
    command: str,
    work_dir: Path,
    artifact_dir: Path,
    output_path: Path,
    timeout: int,
    activity: List[Dict[str, Any]],
) -> Dict[str, Any]:
    started = time.time()
    env = dict(os.environ)
    env.setdefault("PYTHONIOENCODING", "utf-8")
    env["METIS_REPORT_ARTIFACTS_DIR"] = str(artifact_dir)
    env["METIS_REPORT_OUTPUT_PATH"] = str(output_path)
    try:
        if script is not None:
            proc = subprocess.run(
                [sys.executable, str(script)],
                cwd=str(work_dir),
                env=env,
                capture_output=True,
                text=True,
                timeout=timeout,
                check=False,
            )
            command_display = f"{sys.executable} {script}"
        else:
            proc = subprocess.run(
                str(command),
                cwd=str(work_dir),
                env=env,
                capture_output=True,
                text=True,
                timeout=timeout,
                shell=True,
                check=False,
            )
            command_display = str(command)
        duration_ms = int((time.time() - started) * 1000)
        ok = proc.returncode == 0
        _add_activity(
            activity,
            "run_code",
            "Run code",
            ok,
            detail=f"exit {proc.returncode} in {duration_ms}ms",
            command=command_display,
            duration_ms=duration_ms,
        )
        return {
            "ok": ok,
            "returncode": proc.returncode,
            "stdout": proc.stdout or "",
            "stderr": proc.stderr or "",
            "duration_ms": duration_ms,
            "timed_out": False,
        }
    except subprocess.TimeoutExpired as exc:
        duration_ms = int((time.time() - started) * 1000)
        detail = f"timeout after {timeout}s"
        _add_activity(activity, "run_code", "Run code", False, detail=detail, duration_ms=duration_ms)
        return {
            "ok": False,
            "returncode": None,
            "stdout": exc.stdout or "",
            "stderr": exc.stderr or detail,
            "duration_ms": duration_ms,
            "timed_out": True,
        }


def _write_docx_report(
    *,
    target: Path,
    title: str,
    assignment: str,
    source_code: str,
    command_display: str,
    run_result: Dict[str, Any],
    artifacts: List[Dict[str, Any]],
    conclusion: str,
) -> None:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"Missing DOCX dependency: python-docx is required ({type(exc).__name__}: {exc})")

    target.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)

    doc.add_heading(title, level=0)
    if assignment:
        doc.add_heading("Assignment", level=1)
        _add_paragraphs(doc, assignment)

    doc.add_heading("Execution Summary", level=1)
    summary = [
        f"Command: {command_display or 'not provided'}",
        f"Return code: {run_result.get('returncode')}",
        f"Timed out: {bool(run_result.get('timed_out'))}",
        f"Duration: {run_result.get('duration_ms', 0)} ms",
    ]
    for item in summary:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Program Output", level=1)
    _add_preformatted(doc, "stdout", _truncate(run_result.get("stdout", ""), MAX_CAPTURE_CHARS))
    stderr = _truncate(run_result.get("stderr", ""), MAX_CAPTURE_CHARS)
    if stderr:
        _add_preformatted(doc, "stderr", stderr)

    if source_code:
        doc.add_heading("Code", level=1)
        _add_preformatted(doc, "source", _truncate(source_code, MAX_CODE_CHARS))

    if artifacts:
        doc.add_heading("Generated Artifacts", level=1)
        for artifact in artifacts:
            label = f"{artifact.get('kind', 'file')}: {artifact.get('path')} ({artifact.get('size', 0)} bytes)"
            doc.add_paragraph(label, style="List Bullet")

        image_paths = [Path(str(item["path"])) for item in artifacts if item.get("kind") == "image"]
        if image_paths:
            doc.add_heading("Figures", level=1)
            for image in image_paths[:12]:
                doc.add_paragraph(image.name)
                try:
                    doc.add_picture(str(image), width=Inches(5.8))
                except Exception as exc:
                    doc.add_paragraph(f"[Image insert failed: {type(exc).__name__}: {exc}]")

    if conclusion:
        doc.add_heading("Conclusion", level=1)
        _add_paragraphs(doc, conclusion)

    doc.save(str(target))


def _add_preformatted(doc: Any, label: str, text: str) -> None:
    doc.add_paragraph(label)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text or "[empty]")
    run.font.name = "Consolas"
    run.font.size = 9


def _add_paragraphs(doc: Any, text: str) -> None:
    for raw in str(text or "").splitlines():
        if raw.strip():
            doc.add_paragraph(raw.strip())


def _working_dir(raw: str) -> Path:
    if str(raw or "").strip():
        path = safe_path_for_write(str(raw).strip())
        path.mkdir(parents=True, exist_ok=True)
        return path
    return get_workspace_root()


def _artifact_dir(raw: str, target: Path) -> Path:
    if str(raw or "").strip():
        path = safe_path_for_write(str(raw).strip())
    else:
        path = safe_path_for_write(str(Path("output") / "report_artifacts" / target.stem))
    path.mkdir(parents=True, exist_ok=True)
    return path


def _safe_write_path(path: str, label: str) -> Path:
    if not str(path or "").strip():
        raise PathSecurityError(f"{label} is required")
    return safe_path_for_write(str(path))


def _snapshot_files(path: Path) -> set[str]:
    if not path.exists():
        return set()
    return {str(item.resolve(strict=False)) for item in path.rglob("*") if item.is_file()}


def _collect_artifacts(path: Path, before: set[str]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for item in sorted(path.rglob("*")):
        if not item.is_file():
            continue
        resolved = str(item.resolve(strict=False))
        suffix = item.suffix.lower()
        kind = "image" if suffix in IMAGE_EXTENSIONS else "text" if suffix in TEXT_EXTENSIONS else "file"
        generated = resolved not in before
        out.append(
            {
                "path": resolved,
                "name": item.name,
                "kind": kind,
                "size": item.stat().st_size,
                "generated": generated,
            }
        )
    return out[:80]


def _add_activity(
    activity: List[Dict[str, Any]],
    event: str,
    title: str,
    ok: bool,
    *,
    detail: str = "",
    path: str = "",
    command: str = "",
    duration_ms: int | None = None,
) -> None:
    row: Dict[str, Any] = {
        "at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "event": event,
        "title": title,
        "ok": bool(ok),
        "detail": _truncate(detail, 220),
    }
    if path:
        row["path"] = path
    if command:
        row["command"] = command
    if duration_ms is not None:
        row["duration_ms"] = int(duration_ms)
    activity.append(row)


def _activity_summary(status: str, target: Path | None, artifacts: List[Dict[str, Any]]) -> str:
    target_text = str(target) if target else "no report"
    return f"{status}: {target_text}; artifacts={len(artifacts)}"


def _verification_detail(payload: Dict[str, Any]) -> str:
    if not payload:
        return "not inspected"
    if payload.get("render"):
        render = payload.get("render") if isinstance(payload.get("render"), dict) else {}
        images = render.get("images", []) if isinstance(render, dict) else []
        return f"rendered {len(images)} page image(s)" if images else str(render.get("error") or "render unavailable")
    if payload.get("ok"):
        return f"{payload.get('paragraphs', 0)} paragraph(s), {payload.get('tables', 0)} table(s)"
    return str(payload.get("error") or "inspect failed")


def _render_ok(payload: Dict[str, Any]) -> bool:
    render = payload.get("render") if isinstance(payload, dict) else None
    return bool(isinstance(render, dict) and render.get("ok"))


def _truncate(value: Any, limit: int = MAX_CAPTURE_CHARS) -> str:
    text = str(value or "")
    if len(text) <= limit:
        return text
    head = max(0, limit - 80)
    return f"{text[:head]}\n... [truncated {len(text)} chars] ..."


def _json(payload: Dict[str, Any]) -> str:
    return json.dumps(payload, ensure_ascii=False, indent=2)


def _json_error(message: str, **extra: Any) -> str:
    payload: Dict[str, Any] = {"ok": False, "error": message}
    payload.update(extra)
    return _json(payload)
