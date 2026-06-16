from __future__ import annotations

import json
import shutil
import subprocess
import zipfile
from pathlib import Path
from typing import Any, Dict, List
from xml.etree import ElementTree

from backend.tools.coding.foundation.core_mechanisms.path_security import (
    PathSecurityError,
    safe_path_for_read,
    safe_path_for_write,
)

from .pdf_tools import pdf_render_pages


def docx_create(
    output_path: str,
    title: str = "",
    body: str = "",
    sections: List[Dict[str, Any]] | str | None = None,
) -> str:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
    except Exception as exc:
        return _json_error(
            "Missing DOCX dependency: python-docx is required.",
            dependency="python-docx",
            detail=f"{type(exc).__name__}: {exc}",
        )
    target = _resolve_output_path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    if title:
        heading = doc.add_paragraph()
        run = heading.add_run(str(title))
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(20)
    for paragraph in str(body or "").splitlines():
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    for item in _normalize_sections(sections):
        heading = str(item.get("heading") or "").strip()
        text = str(item.get("text") or item.get("body") or "").strip()
        bullets = item.get("bullets") or []
        if heading:
            doc.add_heading(heading, level=max(1, min(int(item.get("level") or 1), 3)))
        if text:
            for paragraph in text.splitlines():
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        if isinstance(bullets, list):
            for bullet in bullets:
                doc.add_paragraph(str(bullet), style="List Bullet")
    doc.save(str(target))
    return _json({"ok": True, "output_path": str(target), "title": title})


def docx_edit(
    path: str,
    output_path: str = "",
    find: str = "",
    replace: str = "",
    append_text: str = "",
) -> str:
    source, error = _resolve_existing_file(path)
    if source is None:
        return _json_error(error or "DOCX file not found", path=path)
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        return _json_error(
            "Missing DOCX dependency: python-docx is required.",
            dependency="python-docx",
            detail=f"{type(exc).__name__}: {exc}",
        )
    doc = Document(str(source))
    replacements = 0
    if find:
        for paragraph in doc.paragraphs:
            if find in paragraph.text:
                _replace_paragraph_text(paragraph, find, replace)
                replacements += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if find in paragraph.text:
                            _replace_paragraph_text(paragraph, find, replace)
                            replacements += 1
    if append_text:
        for paragraph in str(append_text).splitlines():
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    target = _resolve_output_path(output_path) if output_path else source
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    return _json(
        {
            "ok": True,
            "path": str(source),
            "output_path": str(target),
            "replacements": replacements,
            "appended": bool(append_text),
        }
    )


def docx_to_pdf(path: str, output_dir: str = "") -> str:
    source, error = _resolve_existing_file(path)
    if source is None:
        return _json_error(error or "DOCX file not found", path=path)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return _json_error(
            "Missing DOCX renderer: LibreOffice/soffice was not found on PATH.",
            path=str(source),
            dependency="LibreOffice soffice",
            install_hint="Install LibreOffice or bundle soffice with Metis.",
        )
    out_dir = _output_dir(output_dir, "output/docx")
    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(out_dir),
        str(source),
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120, check=False)
    pdf = out_dir / f"{source.stem}.pdf"
    return _json(
        {
            "ok": proc.returncode == 0 and pdf.is_file(),
            "path": str(source),
            "pdf_path": str(pdf),
            "command": cmd,
            "stdout": (proc.stdout or "")[:2000],
            "stderr": (proc.stderr or "")[:2000],
        }
    )


def docx_render_pages(path: str, output_dir: str = "", dpi: int = 150) -> str:
    source, error = _resolve_existing_file(path)
    if source is None:
        return _json_error(error or "DOCX file not found", path=path)
    out_dir = _output_dir(output_dir, "output/docx")
    converted = json.loads(docx_to_pdf(str(source), output_dir=str(out_dir)))
    if not converted.get("ok"):
        return _json(converted)
    rendered = json.loads(pdf_render_pages(str(converted.get("pdf_path")), output_dir=str(out_dir), dpi=dpi))
    return _json(
        {
            "ok": bool(rendered.get("ok")),
            "path": str(source),
            "pdf_path": converted.get("pdf_path", ""),
            "output_dir": str(out_dir),
            "images": rendered.get("images", []),
            "render": rendered,
        }
    )


def docx_inspect_layout(path: str, render: bool = False, output_dir: str = "") -> str:
    source, error = _resolve_existing_file(path)
    if source is None:
        return _json_error(error or "DOCX file not found", path=path)
    payload: Dict[str, Any] = {
        "ok": True,
        "path": str(source),
        "paragraphs": 0,
        "tables": 0,
        "images": 0,
        "headings": [],
    }
    try:
        from docx import Document  # type: ignore

        doc = Document(str(source))
        payload["paragraphs"] = len(doc.paragraphs)
        payload["tables"] = len(doc.tables)
        payload["headings"] = [
            paragraph.text
            for paragraph in doc.paragraphs
            if str(getattr(paragraph.style, "name", "") or "").lower().startswith("heading")
        ][:40]
    except Exception as exc:
        payload["python_docx_error"] = f"{type(exc).__name__}: {exc}"
    payload["images"] = _count_docx_images(source)
    if render:
        payload["render"] = json.loads(docx_render_pages(str(source), output_dir=output_dir))
    return _json(payload)


def _replace_paragraph_text(paragraph: Any, find: str, replace: str) -> None:
    text = paragraph.text.replace(find, replace)
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _normalize_sections(sections: List[Dict[str, Any]] | str | None) -> List[Dict[str, Any]]:
    if sections is None or sections == "":
        return []
    if isinstance(sections, list):
        return [item for item in sections if isinstance(item, dict)]
    try:
        data = json.loads(str(sections))
    except json.JSONDecodeError:
        return [{"heading": "", "text": str(sections)}]
    if isinstance(data, list):
        return [item for item in data if isinstance(item, dict)]
    if isinstance(data, dict):
        return [data]
    return []


def _count_docx_images(path: Path) -> int:
    try:
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
            media = [name for name in names if name.startswith("word/media/")]
            if media:
                return len(media)
            rels = archive.read("word/_rels/document.xml.rels")
            root = ElementTree.fromstring(rels)
            return sum(1 for elem in root.iter() if "image" in str(elem.attrib.get("Type", "")))
    except Exception:
        return 0


def _resolve_existing_file(path: str) -> tuple[Path | None, str]:
    try:
        resolved = safe_path_for_read(str(path or ""))
    except PathSecurityError as exc:
        return None, str(exc)
    return (resolved, "") if resolved.is_file() else (None, "DOCX file not found")


def _resolve_output_path(path: str) -> Path:
    return safe_path_for_write(str(path or ""))


def _output_dir(raw: str, default: str) -> Path:
    path = safe_path_for_write(str(raw or default))
    path.mkdir(parents=True, exist_ok=True)
    return path


def _json(payload: Dict[str, Any]) -> str:
    return json.dumps(payload, ensure_ascii=False, indent=2)


def _json_error(message: str, **extra: Any) -> str:
    payload = {"ok": False, "error": message}
    payload.update(extra)
    return _json(payload)
