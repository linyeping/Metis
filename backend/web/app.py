from __future__ import annotations

import ipaddress
import fnmatch
import hashlib
import json
import logging
import mimetypes
import os
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import time
import uuid
import webbrowser
from dataclasses import replace
from logging.handlers import RotatingFileHandler
from typing import Any, Dict, List, Optional
from urllib.parse import quote, urlparse

from flask import Flask, Response, abort, jsonify, request, send_file, stream_with_context

_WEB_DIR = os.path.dirname(os.path.abspath(__file__))
_BACKEND_ROOT = os.path.dirname(_WEB_DIR)
_REPO_ROOT = os.path.dirname(_BACKEND_ROOT)
if _REPO_ROOT not in sys.path:
    sys.path.insert(0, _REPO_ROOT)

from backend.runtime.agent_loop import (  # noqa: E402
    AgentConfig,
    CompactEvent,
    ContentEvent,
    DoneEvent,
    ErrorEvent,
    PermissionRequestEvent,
    TextDeltaEvent,
    ThinkingEvent,
    ToolCallEvent,
    ToolResultEvent,
    run,
    run_stream,
)
from backend.runtime.cancellation import OperationCancelled  # noqa: E402
from backend.runtime.checkpoints import (  # noqa: E402
    CheckpointRecorder,
    create_checkpoint,
    find_checkpoint,
    list_checkpoints,
    load_compact_state_snapshot,
    load_history_snapshot,
    prune_checkpoints,
    restore_files_from_checkpoint,
)
from backend.core.paths import legacy_miro_path, metis_dir, metis_path  # noqa: E402
from backend.runtime.tool_registry import get_registry  # noqa: E402
from backend.runtime.error_catalog import ErrorInfo, classify_llm_error  # noqa: E402
from backend.runtime.path_safety import validate_path_access, validate_tool_paths  # noqa: E402
from backend.web.desktop_window import handle_window_action  # noqa: E402
from backend.web.error_handlers import register_error_handlers  # noqa: E402
from backend.web.llm_state import (  # noqa: E402
    build_agent_config,
    compaction_stage,
    env,
    env_any,
    env_bool,
    env_disabled,
    first_run_status_payload,
    get_provider_models,
    get_provider_status,
    get_provider_usage,
    get_runtime_settings,
    save_first_run_config,
    sanitize_for_log,
    should_auto_compact,
    update_runtime_settings,
    verify_provider_settings,
)
from backend.web.sessions import get_session_manager  # noqa: E402
from backend.web.workspaces import get_workspace_manager  # noqa: E402
from backend.web.runtime_state import RuntimeState  # noqa: E402
try:
    from backend.bridges.event_contract import agent_event_contract_payload  # noqa: E402
    from backend.bridges.event_serializer import agent_event_payload, sse_data  # noqa: E402
except ImportError:  # pragma: no cover - supports running from inside miro/
    from backend.bridges.event_contract import agent_event_contract_payload  # noqa: E402
    from backend.bridges.event_serializer import agent_event_payload, sse_data  # noqa: E402


def _configure_logging() -> None:
    level_name = os.environ.get("METIS_LOG_LEVEL", "INFO").upper()
    level = getattr(logging, level_name, logging.INFO)
    formatter = logging.Formatter("%(asctime)s [%(levelname)s] %(name)s: %(message)s")
    log_dir = str(metis_dir("logs"))
    log_path = str(metis_path("logs", "metis-backend.log"))

    root = logging.getLogger()
    root.setLevel(level)

    if not any(getattr(handler, "_metis_stream_handler", False) for handler in root.handlers):
        stream_handler = logging.StreamHandler()
        stream_handler.setFormatter(formatter)
        stream_handler.setLevel(level)
        stream_handler._metis_stream_handler = True  # type: ignore[attr-defined]
        root.addHandler(stream_handler)

    if not any(getattr(handler, "_metis_file_handler", False) for handler in root.handlers):
        file_handler = RotatingFileHandler(
            log_path,
            maxBytes=10 * 1024 * 1024,
            backupCount=3,
            encoding="utf-8",
        )
        file_handler.setFormatter(formatter)
        file_handler.setLevel(level)
        file_handler._metis_file_handler = True  # type: ignore[attr-defined]
        root.addHandler(file_handler)

    logging.getLogger("werkzeug").setLevel(level)


_configure_logging()

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024
register_error_handlers(app)
logger = logging.getLogger(__name__)


@app.before_request
def _log_request_start() -> None:
    request.environ["metis_request_started_at"] = time.time()


@app.after_request
def _cors(resp: Response) -> Response:
    origin = request.headers.get("Origin", "")
    if origin and _is_local_origin(origin):
        resp.headers["Access-Control-Allow-Origin"] = origin
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
    resp.headers["Access-Control-Allow-Methods"] = "GET,POST,DELETE,OPTIONS"
    started = float(request.environ.get("metis_request_started_at") or time.time())
    duration_ms = (time.time() - started) * 1000
    log = logger.debug if request.path == "/health" else logger.info
    log("http %s %s -> %s %.1fms", request.method, request.path, resp.status_code, duration_ms)
    return resp


def _is_local_origin(origin: str) -> bool:
    parsed = urlparse(origin)
    return parsed.hostname in {"localhost", "127.0.0.1", "::1"}


_FRONTEND_INDEX = os.path.join(_REPO_ROOT, "desktop", "dist", "index.html")
_ASSETS_DIR = os.path.join(_BACKEND_ROOT, "assets")
_runtime_state = RuntimeState()

from backend.web import helpers as _helpers  # noqa: E402
_helpers.init_shared_state(_runtime_state)

_perm_dict_lock = threading.Lock()
_permission_locks: Dict[str, threading.Event] = {}
_permission_results: Dict[str, bool] = {}
_permission_contexts: Dict[str, Dict[str, Any]] = {}
_URL_TRAILING_MARKERS = ("，", "。", "！", "？", "；", "：", "、", "）", "】", "》", "」", "』")
_SUBAGENT_TOOLS = {
    "task_dispatch",
    "run_parallel_tasks",
    "run_task_graph",
    "delegate_explore",
    "delegate_shell",
    "delegate_browser",
    "delegate_best_of_n",
    "summon_context_gatherer",
}
_COMPOSER_PERMISSION_SOURCE = "composer_access"
_CROSS_WORKSPACE_READ_TOOLS = {
    "read_file",
    "read_multiple_files",
    "list_directory",
    "search_in_file",
    "search_in_codebase",
    "find_files",
    "grep_search",
    "glob_search",
    "semantic_search",
}
_RUN_TERMINAL_STATES = {"done", "failed", "canceled"}
_RUN_ACTIVE_STATES = {"queued", "running", "canceling"}
_RUN_RETENTION_LIMIT = 80
_RUN_DEFAULT_MAX_ACTIVE = 4
_SIDE_CHAT_MAX_MESSAGES = 40
_SIDE_CHAT_MAX_MESSAGE_CHARS = 12000
_SIDE_CHAT_SYSTEM_PROMPT = (
    "You are Metis Chat, a standalone side conversation. "
    "You cannot see the coding agent's private context, workspace files, tools, "
    "run state, or session history unless the user explicitly pastes that information. "
    "Answer normally and do not claim to have run tools."
)
_runs: Dict[str, Dict[str, Any]] = {}
_runs_lock = threading.Lock()


def _workspace_root_for_id(workspace_id: str) -> str:
    workspace = get_workspace_manager().get_workspace(workspace_id) if workspace_id else None
    if workspace and workspace.path:
        return os.path.abspath(workspace.path)
    return ""


def _workspace_root_for_session(session_id: str) -> str:
    session = get_session_manager().get_session(session_id) if session_id else None
    if session and session.workspace_id:
        return _workspace_root_for_id(session.workspace_id)
    return ""


def _active_workspace_root() -> str:
    return _workspace_root_for_id(_runtime_state.active_workspace_id) or os.path.abspath(os.getcwd())


def _request_workspace_root(session_id: str = "") -> str:
    return _workspace_root_for_session(session_id) or _active_workspace_root()


def _load_config(workspace_root: str = "") -> AgentConfig:
    root = os.path.abspath(workspace_root or _active_workspace_root())
    return build_agent_config(
        system_prompt=_load_system_prompt(),
        user_memory_text=_load_metis_md(root),
        execution_mode=_runtime_state.execution_mode,
        permission_checker=_permission_checker_for_workspace(root),
        tool_boundary_overrides=_tool_boundary_overrides_for_workspace(root),
        workspace_root=root,
    )


def _load_config_for_workspace(workspace_root: str = "") -> AgentConfig:
    """Load config while tolerating older no-arg test monkeypatches."""
    root = os.path.abspath(workspace_root or _active_workspace_root())
    try:
        config = _load_config(root)
    except TypeError:
        config = _load_config()
        config.workspace_root = root
    return config


def _load_system_prompt() -> str:
    prompt_path = os.path.join(_BACKEND_ROOT, "core", "prompts", "MAIN_PROMPT.txt")
    default_md_path = os.path.join(_BACKEND_ROOT, "core", "prompts", "METIS_DEFAULT.md")
    parts: List[str] = []
    if os.path.exists(prompt_path):
        with open(prompt_path, "r", encoding="utf-8") as handle:
            parts.append(handle.read())
    else:
        parts.append(
            "You are Metis, an AI assistant with coding and desktop automation "
            "capabilities. Use tools when they help answer the user."
        )
    if os.path.exists(default_md_path):
        with open(default_md_path, "r", encoding="utf-8") as handle:
            parts.append(handle.read())
    return "\n\n---\n\n".join(part.strip() for part in parts if part.strip())


def _load_metis_md(workspace_root: str = "") -> str:
    """Load global and project METIS.md memory, with MIRO.md fallback."""
    root = workspace_root or _active_workspace_root()
    sections: List[str] = []
    memory_pairs = [
        (
            f"Global Memory - {metis_path('METIS.md')}",
            str(metis_path("METIS.md")),
            f"Global Legacy Memory - {legacy_miro_path('MIRO.md')}",
            str(legacy_miro_path("MIRO.md")),
        ),
        (
            "Project Memory - METIS.md",
            os.path.join(root, "METIS.md"),
            "Project Legacy Memory - MIRO.md",
            os.path.join(root, "MIRO.md"),
        ),
    ]

    for primary_label, primary_path, legacy_label, legacy_path in memory_pairs:
        label = primary_label
        path = primary_path
        if not os.path.isfile(path) and os.path.isfile(legacy_path):
            label = legacy_label
            path = legacy_path
        if not os.path.isfile(path):
            continue
        try:
            with open(path, "r", encoding="utf-8") as handle:
                content = handle.read().strip()
        except OSError:
            continue
        if content:
            sections.append(f"[{label}]\n{content}")
    return "\n\n".join(sections)


def _build_full_system_prompt(workspace_root: str = "", model: str = "") -> str:
    """Compile the effective system prompt for diagnostics and verification."""
    from backend.bridges.model_capability import detect_from_model_name
    from backend.core.engine.prompt_runtime import compile_prompt_runtime

    capabilities = detect_from_model_name(model)
    snapshot = compile_prompt_runtime(
        _load_system_prompt(),
        user_memory_text=_load_metis_md(workspace_root),
        model_tier=capabilities.tier,
        model_context_window=capabilities.effective_context,
        workspace_root=workspace_root or "",
    )
    return snapshot.final_system_prompt


def _estimate_prompt_tokens(messages: List[Dict[str, Any]], system_prompt: str = "") -> int:
    total = len(system_prompt or "")
    for message in messages:
        total += len(_export_content(message.get("content", "")))
        total += len(str(message.get("name") or ""))
        tool_calls = message.get("tool_calls") or []
        if tool_calls:
            total += len(json.dumps(tool_calls, ensure_ascii=False, default=str))
    return total // 4


def _strip_images_from_content(content: Any) -> Any:
    if isinstance(content, list):
        stripped = []
        for block in content:
            if not isinstance(block, dict):
                stripped.append(block)
                continue
            block_type = str(block.get("type") or "").lower()
            if "image" in block_type:
                continue
            stripped.append(block)
        return stripped
    return content


def _mask_observations(
    history: List[Dict[str, Any]],
    keep_recent: int = 6,
    *,
    aggressive: bool = False,
) -> List[Dict[str, Any]]:
    if keep_recent <= 0:
        old = list(history)
        recent: List[Dict[str, Any]] = []
    elif len(history) <= keep_recent:
        return list(history)
    else:
        old = history[:-keep_recent]
        recent = history[-keep_recent:]
    masked: List[Dict[str, Any]] = []

    for message in old:
        role = message.get("role")
        content = _export_content(message.get("content", ""))
        if role == "tool":
            tool_name = message.get("name", "tool")
            preview_limit = 200 if aggressive else 100
            preview = content[:preview_limit].replace("\n", " ")
            masked.append(
                {
                    "role": "tool",
                    "name": tool_name,
                    "tool_call_id": message.get("tool_call_id", ""),
                    "content": f"[Observation masked — {tool_name}: {preview}...]",
                }
            )
            continue

        next_message = dict(message)
        if role == "assistant":
            if aggressive and len(content) > 600:
                next_message["content"] = content[:600] + "\n[...truncated]"
            masked.append(next_message)
            continue

        if aggressive:
            next_message["content"] = _strip_images_from_content(next_message.get("content", ""))
        truncated_limit = 300 if aggressive else 500
        if len(content) > truncated_limit:
            next_message["content"] = content[:truncated_limit] + "\n[...truncated]"
        masked.append(next_message)

    recent_messages = []
    for message in recent:
        next_message = dict(message)
        if aggressive:
            next_message["content"] = _strip_images_from_content(next_message.get("content", ""))
            if next_message.get("role") == "tool":
                content = _export_content(next_message.get("content", ""))
                next_message["content"] = content[:200] + ("..." if len(content) > 200 else "")
        recent_messages.append(next_message)

    return masked + recent_messages


def _permission_paths(workspace_root: str = "") -> Dict[str, str]:
    root = os.path.abspath(workspace_root or _active_workspace_root())
    metis_dir = os.path.join(root, ".metis")
    return {
        "config": os.path.join(metis_dir, "permissions.json"),
        "legacy_config": os.path.join(root, ".miro", "permissions.json"),
        "audit": os.path.join(metis_dir, "audit", "tool-permissions.jsonl"),
    }


def _load_permission_document(workspace_root: str = "") -> Dict[str, Any]:
    """Load workspace permission rules from .metis, with .miro read fallback."""
    paths = _permission_paths(workspace_root)
    source_path = paths["config"]
    if not os.path.isfile(source_path) and os.path.isfile(paths["legacy_config"]):
        source_path = paths["legacy_config"]
    if not os.path.isfile(source_path):
        return {"rules": [], "source_path": paths["config"], "exists": False}
    try:
        with open(source_path, "r", encoding="utf-8") as handle:
            data = json.load(handle)
    except (json.JSONDecodeError, OSError):
        return {"rules": [], "source_path": source_path, "exists": True}
    rules = data.get("rules", []) if isinstance(data, dict) else []
    if not isinstance(rules, list):
        rules = []
    return {"rules": rules, "source_path": source_path, "exists": True}


def _normalize_permission_rule(rule: Dict[str, Any]) -> Dict[str, Any]:
    now = time.time()
    action = str(rule.get("action") or "ask").strip().lower()
    if action not in {"allow", "deny", "ask"}:
        action = "ask"
    args_match = rule.get("args_match", {})
    if not isinstance(args_match, dict):
        args_match = {}
    return {
        "id": str(rule.get("id") or uuid.uuid4()),
        "tool": str(rule.get("tool") or "*").strip() or "*",
        "action": action,
        "args_match": {
            str(key): str(value)
            for key, value in args_match.items()
            if str(key).strip() and str(value).strip()
        },
        "source": str(rule.get("source") or "workspace").strip() or "workspace",
        "created_at": float(rule.get("created_at") or now),
        "updated_at": float(rule.get("updated_at") or rule.get("created_at") or now),
    }


def _permission_rules(workspace_root: str = "") -> List[Dict[str, Any]]:
    document = _load_permission_document(workspace_root)
    return [
        _normalize_permission_rule(rule)
        for rule in document.get("rules", [])
        if isinstance(rule, dict)
    ]


def _write_permission_rules(rules: List[Dict[str, Any]], workspace_root: str = "") -> None:
    path = _permission_paths(workspace_root)["config"]
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as handle:
        json.dump({"rules": rules}, handle, ensure_ascii=False, indent=2)


def _load_permission_rules(workspace_root: str = "") -> List[Dict[str, Any]]:
    return _permission_rules(workspace_root)


def _permission_rule_payload(rule: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "id": str(rule.get("id") or ""),
        "tool": str(rule.get("tool") or "*"),
        "action": str(rule.get("action") or "ask"),
        "args_match": dict(rule.get("args_match") or {}),
        "source": str(rule.get("source") or "workspace"),
        "created_at": float(rule.get("created_at") or 0),
        "updated_at": float(rule.get("updated_at") or 0),
    }


def _create_permission_rule(
    *,
    tool: str,
    action: str,
    args_match: Optional[Dict[str, Any]] = None,
    source: str = "workspace",
    workspace_root: str = "",
) -> Dict[str, Any]:
    rule = _normalize_permission_rule(
        {
            "tool": tool,
            "action": action,
            "args_match": args_match or {},
            "source": source,
        }
    )
    rules = _permission_rules(workspace_root)
    for index, existing in enumerate(rules):
        if (
            existing.get("tool") == rule["tool"]
            and existing.get("action") == rule["action"]
            and existing.get("args_match") == rule["args_match"]
            and existing.get("source") == rule["source"]
        ):
            existing["updated_at"] = time.time()
            rules[index] = existing
            _write_permission_rules(rules, workspace_root)
            return existing
    rules.append(rule)
    _write_permission_rules(rules, workspace_root)
    return rule


def _sanitize_permission_args(value: Any, depth: int = 0) -> Any:
    if depth > 4:
        return "[truncated]"
    if isinstance(value, dict):
        sanitized: Dict[str, Any] = {}
        for key, item in list(value.items())[:40]:
            key_text = str(key)
            lower = key_text.lower()
            if any(marker in lower for marker in ("api_key", "apikey", "token", "secret", "password", "authorization")):
                sanitized[key_text] = "***"
            else:
                sanitized[key_text] = _sanitize_permission_args(item, depth + 1)
        return sanitized
    if isinstance(value, list):
        return [_sanitize_permission_args(item, depth + 1) for item in value[:20]]
    if isinstance(value, str):
        return value if len(value) <= 220 else value[:217] + "..."
    return value


def _append_permission_audit(entry: Dict[str, Any], workspace_root: str = "") -> Dict[str, Any]:
    root = os.path.abspath(workspace_root or str(entry.get("workspace_root") or "") or _active_workspace_root())
    paths = _permission_paths(root)
    audit = {
        "id": str(entry.get("id") or uuid.uuid4()),
        "created_at": float(entry.get("created_at") or time.time()),
        "workspace_id": _runtime_state.active_workspace_id or "",
        "session_id": _runtime_state.active_session_id or "",
        "cwd": root,
        "request_id": str(entry.get("request_id") or ""),
        "call_id": str(entry.get("call_id") or ""),
        "tool": str(entry.get("tool") or ""),
        "action": str(entry.get("action") or ""),
        "approved": bool(entry.get("approved", False)),
        "remember": str(entry.get("remember") or ""),
        "rule_id": str(entry.get("rule_id") or ""),
        "source": str(entry.get("source") or "permission_dialog"),
        "arguments": _sanitize_permission_args(entry.get("arguments") or {}),
    }
    os.makedirs(os.path.dirname(paths["audit"]), exist_ok=True)
    with open(paths["audit"], "a", encoding="utf-8") as handle:
        handle.write(json.dumps(audit, ensure_ascii=False, sort_keys=True) + "\n")
    return audit


def _read_permission_audit(limit: int = 50, workspace_root: str = "") -> List[Dict[str, Any]]:
    path = _permission_paths(workspace_root)["audit"]
    if not os.path.isfile(path):
        return []
    try:
        with open(path, "r", encoding="utf-8") as handle:
            lines = handle.readlines()
    except OSError:
        return []
    entries: List[Dict[str, Any]] = []
    for line in reversed(lines[-max(1, limit * 2) :]):
        try:
            row = json.loads(line)
        except json.JSONDecodeError:
            continue
        if isinstance(row, dict):
            entries.append(row)
        if len(entries) >= limit:
            break
    return entries


def _permission_rule_matches(rule: Dict[str, Any], tool_name: str, arguments: Dict[str, Any]) -> bool:
    rule_tool = str(rule.get("tool", ""))
    if rule_tool and rule_tool != "*" and not fnmatch.fnmatch(tool_name, rule_tool):
        return False
    args_match = rule.get("args_match", {})
    if isinstance(args_match, dict) and args_match:
        for key, pattern in args_match.items():
            value = str(arguments.get(str(key), ""))
            if not fnmatch.fnmatch(value, str(pattern)):
                return False
    return True


def _is_composer_full_access_rule(rule: Dict[str, Any]) -> bool:
    return (
        str(rule.get("source") or "") == _COMPOSER_PERMISSION_SOURCE
        and str(rule.get("tool") or "") == "*"
        and str(rule.get("action") or "").lower() == "allow"
        and not dict(rule.get("args_match") or {})
    )


def _composer_full_access_enabled(workspace_root: str = "") -> bool:
    return any(_is_composer_full_access_rule(rule) for rule in _load_permission_rules(workspace_root))


def _tool_boundary_overrides(tool_name: str, arguments: Dict[str, Any], workspace_root: str = "") -> Dict[str, bool]:
    if tool_name not in _CROSS_WORKSPACE_READ_TOOLS:
        return {}
    if not _composer_full_access_enabled(workspace_root):
        return {}
    return {
        "allow_paths_outside_workspace": True,
        "allow_search_outside_workspace": True,
        "allow_semantic_outside_workspace": True,
        "allow_notebook_paths_outside_workspace": True,
    }


def _tool_boundary_overrides_for_workspace(workspace_root: str) -> Any:
    root = os.path.abspath(workspace_root or _active_workspace_root())
    return lambda tool_name, arguments: _tool_boundary_overrides(tool_name, arguments, root)


def _permission_checker_for_workspace(workspace_root: str) -> Any:
    root = os.path.abspath(workspace_root or _active_workspace_root())
    return lambda tool_name, arguments: _check_permission_rules(tool_name, arguments, root)


def _check_permission_rules(tool_name: str, arguments: Dict[str, Any], workspace_root: str = "") -> Optional[str]:
    """Return allow, deny, ask, or None based on .miro/permissions.json."""
    root = os.path.abspath(workspace_root or _active_workspace_root())
    safety = validate_tool_paths(tool_name, arguments, workspace_root=root)
    if not safety.allowed:
        return "deny"

    rules = sorted(
        _load_permission_rules(root),
        key=lambda rule: (
            0 if dict(rule.get("args_match") or {}) else 1,
            0 if str(rule.get("source") or "") == "composer_access" else 1,
            -float(rule.get("updated_at") or rule.get("created_at") or 0),
        ),
    )
    for rule in rules:
        if not isinstance(rule, dict):
            continue
        if not _permission_rule_matches(rule, tool_name, arguments):
            continue
        action = str(rule.get("action", "ask")).lower()
        if action in {"allow", "deny", "ask"}:
            return action
    return None


def _save_active_session() -> None:
    """Persist the current in-memory chat history to the active session."""
    from backend.web.session_routes import save_active_session
    save_active_session()


def _save_session_history(
    session_id: Optional[str],
    *,
    history: List[Dict[str, Any]],
    compact_state: Optional[Dict[str, Any]] = None,
    mode: str = "auto",
) -> bool:
    """Persist chat history to a specific session without relying on global active state."""
    from backend.web.session_routes import save_session_history
    return save_session_history(session_id, history=history, compact_state=compact_state, mode=mode)


def _disk_full_event() -> Dict[str, Any]:
    return _event_payload(
        ErrorEvent(
            recoverable=True,
            code="DISK_FULL",
            title="会话保存失败",
            message="磁盘空间不足或文件系统不可写，会话未保存。",
            hint="请释放磁盘空间后继续；当前内存中的对话仍可暂时查看。",
        )
    )


def _auto_title(message: str) -> str:
    """Generate a short session title from the first user message."""
    text = message.strip()
    if len(text) > 40:
        text = text[:37] + "..."
    return text or "New chat"


def _message_text(message: Any) -> str:
    """Extract readable text from a string or multimodal content blocks."""
    if isinstance(message, str):
        return message
    if isinstance(message, list):
        for block in message:
            if isinstance(block, dict) and block.get("type") == "text":
                return str(block.get("text") or "")
        return ""
    return str(message or "")


def _new_message(role: str, content: Any, **extra: Any) -> Dict[str, Any]:
    message = {"id": f"msg_{uuid.uuid4().hex}", "role": role, "content": content}
    message.update({key: value for key, value in extra.items() if value is not None})
    return message


_TOOL_RECORD_MAX_CHARS = 4000


def _truncate_tool_record(result: Any) -> str:
    """Store a bounded tool result summary in the transcript (avoid bloating the
    session file with screenshots / huge stdout). UI shows this summary."""
    text = result if isinstance(result, str) else str(result or "")
    if len(text) <= _TOOL_RECORD_MAX_CHARS:
        return text
    head = _TOOL_RECORD_MAX_CHARS * 3 // 4
    tail = _TOOL_RECORD_MAX_CHARS // 4
    return f"{text[:head]}\n... [结果已截断，共 {len(text)} 字符] ...\n{text[-tail:]}"


def _tool_result_is_error(result: Any) -> bool:
    text = str(result or "").lstrip()
    return text.startswith(("❌", "Error", "错误", "[Permission denied]", "[Cancelled]")) or (
        "Traceback (most recent call last)" in text
    )


def _safe_int(value: Any, fallback: int = 0) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return fallback


def _safe_float(value: Any, fallback: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return fallback


def _normalize_compact_state(value: Any) -> Dict[str, Any]:
    if not isinstance(value, dict):
        return {}
    summary = str(value.get("summary") or "").strip()
    if not summary:
        return {}
    return {
        "summary": summary,
        "boundary_message_id": str(value.get("boundary_message_id") or ""),
        "boundary_index": max(0, _safe_int(value.get("boundary_index"), 0)),
        "compacted_at": _safe_float(value.get("compacted_at"), 0.0),
        "compact_count": max(0, _safe_int(value.get("compact_count"), 0)),
    }


def _compact_boundary_index(history: List[Dict[str, Any]], compact_state: Any) -> int:
    state = _normalize_compact_state(compact_state)
    if not state:
        return 0
    boundary_id = str(state.get("boundary_message_id") or "")
    if boundary_id:
        for index, message in enumerate(history):
            if isinstance(message, dict) and str(message.get("id") or "") == boundary_id:
                return index
    return min(max(0, _safe_int(state.get("boundary_index"), 0)), len(history))


def _model_context_for_history(history: List[Dict[str, Any]], compact_state: Any) -> List[Dict[str, Any]]:
    state = _normalize_compact_state(compact_state)
    if not state:
        return list(history)
    boundary_index = _compact_boundary_index(history, state)
    return [{"role": "system", "content": state["summary"]}] + list(history[boundary_index:])


def _model_context_with_skill_invocation(
    history: List[Dict[str, Any]],
    compact_state: Any,
    workspace_root: str,
) -> List[Dict[str, Any]]:
    messages = _model_context_for_history(history, compact_state)
    if not messages or str(messages[-1].get("role") or "") != "user":
        return messages
    try:
        from backend.runtime.skill_loader import expand_user_skill_command

        original = _message_text(messages[-1])
        expanded = expand_user_skill_command(original, workspace_root=workspace_root)
        if expanded != original:
            messages[-1] = {**messages[-1], "content": expanded}
    except Exception:
        pass
    return messages


def _compact_state_after_truncate(
    original_history: List[Dict[str, Any]],
    new_history: List[Dict[str, Any]],
    compact_state: Any,
) -> Dict[str, Any]:
    state = _normalize_compact_state(compact_state)
    if not state:
        return {}
    boundary_index = _compact_boundary_index(original_history, state)
    if len(new_history) <= boundary_index:
        return {}
    boundary_id = str(state.get("boundary_message_id") or "")
    if boundary_id and not any(str(message.get("id") or "") == boundary_id for message in new_history if isinstance(message, dict)):
        return {}
    return {
        **state,
        "boundary_index": _compact_boundary_index(new_history, state),
    }


def _compact_summary_from_messages(messages: List[Dict[str, Any]]) -> str:
    if not messages:
        return ""
    first = messages[0] if isinstance(messages[0], dict) else {}
    return _export_content(first.get("content", "")).strip()


def _active_run_for_session(session_id: str) -> Optional[Dict[str, Any]]:
    if not session_id:
        return None
    with _runs_lock:
        active_for_session = [
            run
            for run in _runs.values()
            if run.get("session_id") == session_id and str(run.get("status")) in _RUN_ACTIVE_STATES
        ]
    if not active_for_session:
        return None
    active_for_session.sort(key=lambda run: float(run.get("created_at") or 0), reverse=True)
    return active_for_session[0]


def _sync_active_runtime_history(session_id: Optional[str], history: List[Dict[str, Any]], mode: str) -> None:
    if session_id and session_id == _runtime_state.active_session_id:
        _runtime_state.chat_history = list(history)
        _runtime_state.execution_mode = mode or "auto"


def _compact_status_payload(
    *,
    ok: bool,
    before_context_messages: int,
    after_context_messages: int,
    summary_preview: str = "",
    error: str = "",
    running: bool = False,
) -> Dict[str, Any]:
    return {
        "running": running,
        "ok": ok,
        "before_count": before_context_messages,
        "after_count": after_context_messages,
        "before_context_messages": before_context_messages,
        "after_context_messages": after_context_messages,
        "summary_preview": summary_preview,
        "updated_at": time.time(),
        "error": error,
    }


def _export_content(content: Any) -> str:
    """Convert message content to readable Markdown-safe text for export."""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts: List[str] = []
        for block in content:
            if isinstance(block, str):
                parts.append(block)
            elif isinstance(block, dict):
                if block.get("type") == "text":
                    parts.append(str(block.get("text") or ""))
                elif block.get("type") == "image_url":
                    parts.append("[Image attachment]")
                elif block.get("text") or block.get("content"):
                    parts.append(str(block.get("text") or block.get("content") or ""))
        return "\n".join(part for part in parts if part)
    return str(content or "")


def _extract_key_facts(history: List[Dict[str, Any]]) -> str:
    """Extract structured key facts from history before compaction (write-before-compaction)."""
    changed_files: List[str] = []
    decisions: List[str] = []
    errors: List[str] = []
    for msg in history:
        role = msg.get("role", "")
        content = _export_content(msg.get("content", ""))
        if role == "tool":
            tool_name = msg.get("name", "")
            if any(kw in tool_name for kw in ("write", "edit", "replace", "append", "delete", "create", "patch", "rename")):
                # Extract file path from first line of result
                first_line = content.split("\n", 1)[0][:120]
                changed_files.append(f"  - {tool_name}: {first_line}")
        elif role == "assistant" and content:
            lower = content[:500].lower()
            if any(kw in lower for kw in ("error", "fail", "fix", "bug", "issue")):
                errors.append(f"  - {content[:150]}")
            if any(kw in lower for kw in ("decided", "choosing", "approach", "strategy", "plan")):
                decisions.append(f"  - {content[:150]}")
    parts: List[str] = []
    if changed_files:
        parts.append("Files changed:\n" + "\n".join(changed_files[-15:]))
    if decisions:
        parts.append("Key decisions:\n" + "\n".join(decisions[-5:]))
    if errors:
        parts.append("Errors encountered:\n" + "\n".join(errors[-5:]))
    return "\n".join(parts) if parts else ""


def _mechanical_compact(
    history: List[Dict[str, Any]],
    keep_recent: int = 4,
) -> List[Dict[str, Any]]:
    """No-LLM fallback compaction: drop tool echoes, truncate old messages."""
    if len(history) <= keep_recent:
        return list(history)
    old = history[:-keep_recent]
    recent = history[-keep_recent:]
    key_facts = _extract_key_facts(old)
    compacted: List[Dict[str, Any]] = []
    for msg in old:
        role = msg.get("role", "")
        if role == "tool":
            compacted.append({
                "role": "tool",
                "tool_call_id": msg.get("tool_call_id", ""),
                "name": msg.get("name", ""),
                "content": "[Result omitted during compaction]",
            })
        elif role == "assistant":
            content = _export_content(msg.get("content", ""))
            truncated = content[:300] + "[...]" if len(content) > 300 else content
            compacted.append({**msg, "content": truncated})
        else:
            content = _export_content(msg.get("content", ""))
            truncated = content[:200] + "[...]" if len(content) > 200 else content
            compacted.append({**msg, "content": truncated})
    # Keep only last 6 compacted messages to limit size
    tail_compacted = compacted[-6:] if len(compacted) > 6 else compacted
    summary_content = (
        f"[Mechanical compaction: {len(old)} messages compacted without summarization]\n\n"
        + (key_facts if key_facts else "(no key facts extracted)")
    )
    return [{"role": "system", "content": summary_content}] + tail_compacted + recent


def _compact_history(
    history: List[Dict[str, Any]],
    keep_recent: int = 4,
    *,
    aggressive: bool = False,
) -> Optional[List[Dict[str, Any]]]:
    """Multi-tier summarization with mechanical fallback."""
    if len(history) <= keep_recent + 2:
        return None

    old_messages = history[:-keep_recent]
    recent_messages = history[-keep_recent:]

    # Write-before-compaction: extract key facts first
    key_facts = _extract_key_facts(old_messages)

    # Build summary input with tier-appropriate truncation
    summary_lines: List[str] = []
    source_messages = _mask_observations(old_messages, keep_recent=0, aggressive=aggressive)
    for message in source_messages:
        role = message.get("role", "unknown")
        content = _export_content(message.get("content", ""))
        if role == "tool":
            tool_name = message.get("name", "tool")
            if aggressive:
                summary_lines.append(f"[tool:{tool_name}] {'OK' if len(content) < 100 else content[:80]}")
            else:
                summary_lines.append(f"[tool:{tool_name}] {content[:200]}")
        else:
            limit = 200 if aggressive else 500
            summary_lines.append(f"[{role}] {content[:limit]}")

    summary_input = "\n".join(summary_lines)
    if len(summary_input) > 8000:
        summary_input = summary_input[:8000] + "\n[...truncated...]"

    # Prepend key facts so LLM preserves them in the summary
    if key_facts:
        summary_input = f"KEY FACTS (preserve these):\n{key_facts}\n\n---\nCONVERSATION:\n{summary_input}"

    # Attempt LLM summarization with retry
    summary = ""
    for attempt in range(2):
        try:
            from backend.runtime.agent_loop import _create_backend

            config = _load_config_for_workspace(_active_workspace_root())
            backend = _create_backend(config)
            timeout = 15.0 if attempt > 0 else 30.0
            use_aggressive = aggressive or attempt > 0
            response = backend.chat(
                [
                    {
                        "role": "system",
                        "content": (
                            "You are summarizing conversation history for context continuity. "
                            "The project's METIS.md is always in the system prompt, so do not repeat it.\n\n"
                            + (
                                (
                                    "Produce a single compact paragraph covering decisions, changed files, "
                                    "errors and resolutions, current task, and user preferences. "
                                    "Preserve all KEY FACTS listed at the top. "
                                    "Use the same language as the conversation."
                                )
                                if use_aggressive
                                else (
                                    "Produce a structured summary with these sections (skip empty ones):\n"
                                    "## Decisions Made\n"
                                    "## Files Changed\n"
                                    "## Errors & Resolutions\n"
                                    "## Current Task\n"
                                    "## User Preferences\n\n"
                                    "Preserve all KEY FACTS listed at the top. "
                                    "Use the same language as the conversation. Be concise and use bullets only."
                                )
                            )
                        ),
                    },
                    {"role": "user", "content": summary_input},
                ],
                temperature=0.1,
                max_tokens=1500,
                timeout=timeout,
            )
            summary = response.content.strip() if response.content else ""
            if summary:
                break
        except Exception:
            if attempt == 0:
                continue
            break

    if not summary:
        # Fallback to mechanical compaction — never returns None
        return _mechanical_compact(history, keep_recent=keep_recent)

    summary_title = (
        "[Context Summary - aggressively compacted from "
        if aggressive
        else "[Context Summary - auto-compacted from "
    )

    return [
        {
            "role": "system",
            "content": summary_title + f"{len(old_messages)} earlier messages]\n\n{summary}",
        },
    ] + recent_messages


def _trigger_auto_compact(config: AgentConfig, *, aggressive: bool = False) -> Optional[CompactEvent]:
    """Report runtime-only compaction without rewriting the persisted transcript."""
    model_context = _model_context_for_history(_runtime_state.chat_history, _runtime_state.compact_state)
    before_count = len(model_context)
    if before_count < 5:
        return None

    preview = "本次运行已整理模型上下文；可见对话历史保持完整。"
    _runtime_state.last_compact_status = _compact_status_payload(
        ok=True,
        before_context_messages=before_count,
        after_context_messages=before_count,
        summary_preview=preview,
    )
    return CompactEvent(
        before_count=before_count,
        after_count=before_count,
        summary_preview=preview,
    )


def _generate_fallback_title(text: str) -> str:
    """Generate a local title without LLM, used as fallback."""
    text = text.strip()
    if not text:
        return ""
    first_line = text.split("\n")[0].strip()
    # For CJK-heavy text, use character count
    if sum(1 for c in first_line if "\u4e00" <= c <= "\u9fff") > len(first_line) * 0.3:
        if len(first_line) <= 18:
            return first_line
        return first_line[:16] + "..."
    # For Latin text, truncate at word boundary
    if len(first_line) <= 20:
        return first_line
    words = first_line.split()
    title = ""
    for word in words:
        if len(title) + len(word) + 1 > 18:
            break
        title = f"{title} {word}".strip()
    return (title or first_line[:16]) + "..."


def _generate_smart_title(session_id: str, user_message: Any) -> None:
    """Generate a concise session title using the LLM in a background thread."""
    title_text = _message_text(user_message)
    if not title_text.strip():
        return

    def _run() -> None:
        try:
            config = _load_config_for_workspace(_active_workspace_root())
            from backend.runtime.agent_loop import _create_backend

            backend = _create_backend(config)
            prompt_messages = [
                {
                    "role": "system",
                    "content": (
                        "Generate a very short title (maximum 20 characters) for a chat session "
                        "based on the user's first message. The title should be concise and descriptive. "
                        "Reply with the same language as the user's message. "
                        "Output ONLY the title text, nothing else. No quotes, no punctuation, no explanation."
                    ),
                },
                {"role": "user", "content": title_text},
            ]
            response = backend.chat(prompt_messages, temperature=0.3, max_tokens=30, timeout=30.0)
            title = ""
            if hasattr(response, "content"):
                content = response.content
                if isinstance(content, list):
                    for block in content:
                        if hasattr(block, "text"):
                            title = str(block.text)
                            break
                        if isinstance(block, dict) and block.get("text"):
                            title = str(block["text"])
                            break
                elif isinstance(content, str):
                    title = content
            elif hasattr(response, "choices"):
                title = response.choices[0].message.content or ""
            else:
                title = str(response) if response else ""

            title = title.strip().strip("\"'`").strip()
            title = title.strip(".,:;!?")
            title = title.strip("\u3002\uff0c\uff01\uff1f\u3001\uff1b\uff1a")
            if len(title) > 20:
                title = title[:17] + "..."
            if title:
                get_session_manager().update_session(session_id, title=title)
        except Exception:
            # LLM title failed \u2014 use local fallback instead of silent pass
            fallback = _generate_fallback_title(title_text)
            if fallback:
                try:
                    get_session_manager().update_session(session_id, title=fallback)
                except Exception:
                    pass

    thread = threading.Thread(target=_run, daemon=True, name="miro-auto-title")
    thread.start()



# _ext_to_lang moved to workspace_routes.py


def _init_session() -> None:
    """Load the initial session at startup without creating an empty chat."""
    workspace_manager = get_workspace_manager()
    session_manager = get_session_manager()
    default_workspace = workspace_manager.create_workspace(os.getcwd())
    _runtime_state.active_workspace_id = default_workspace.id
    session_manager.assign_unscoped_sessions(default_workspace.id)
    _load_latest_session_for_workspace(_runtime_state.active_workspace_id)


def _clear_active_session_state() -> None:
    from backend.web.session_routes import clear_active_session_state
    clear_active_session_state()


def _load_latest_session_for_workspace(workspace_id: str) -> Optional[str]:
    from backend.web.session_routes import load_latest_session_for_workspace
    return load_latest_session_for_workspace(workspace_id)


def _ensure_active_session_for_message(user_message: Any) -> None:
    manager = get_session_manager()
    if _runtime_state.active_session_id and manager.get_session(_runtime_state.active_session_id) is not None:
        return
    session = manager.create_session(
        title=_auto_title(_message_text(user_message)),
        workspace_id=_runtime_state.active_workspace_id or "",
    )
    _runtime_state.activate_session(session.id, compact_state=dict(session.compact_state), mode=session.mode)


def _request_session_id() -> str:
    if request.method == "GET":
        return str(request.args.get("session_id") or request.args.get("sessionId") or "").strip()
    data = request.get_json(silent=True) or {}
    return str(data.get("session_id") or data.get("sessionId") or "").strip()


def _prepare_chat_session(user_message: Any) -> tuple[str, List[Dict[str, Any]], Dict[str, Any], str, bool]:
    """Return the session targeted by this chat request and an editable history copy."""
    requested_session_id = _request_session_id()
    manager = get_session_manager()
    if requested_session_id:
        session = manager.get_session(requested_session_id)
        if session is None:
            abort(404, description="session not found")
        if _runtime_state.active_session_id == session.id:
            return session.id, list(_runtime_state.chat_history), dict(_runtime_state.compact_state), _runtime_state.execution_mode, True
        return session.id, list(session.history), dict(session.compact_state), session.mode, True

    _ensure_active_session_for_message(user_message)
    return (
        _runtime_state.active_session_id or "",
        list(_runtime_state.chat_history),
        dict(_runtime_state.compact_state),
        _runtime_state.execution_mode,
        False,
    )


def _commit_request_history_to_active(
    session_id: str,
    history: List[Dict[str, Any]],
    compact_state: Dict[str, Any],
    mode: str,
    explicit_session: bool,
) -> None:
    """Keep the global runtime in sync only when the request owns the active session."""
    if not session_id:
        return
    if explicit_session and _runtime_state.active_session_id != session_id:
        return
    _runtime_state.activate_session(session_id, history=history, compact_state=compact_state, mode=mode)


def _create_run_state(
    *,
    session_id: str,
    assistant_id: str,
    history: List[Dict[str, Any]],
    mode: str,
    model_context: Optional[List[Dict[str, Any]]] = None,
    workspace_root: str = "",
    checkpoint: Optional[CheckpointRecorder] = None,
) -> Dict[str, Any]:
    now = time.time()
    run = {
        "id": uuid.uuid4().hex,
        "session_id": session_id,
        "assistant_id": assistant_id,
        "mode": mode or "auto",
        "workspace_root": os.path.abspath(workspace_root or _request_workspace_root(session_id)),
        "status": "queued",
        "phase": "queued",
        "history": list(history),
        "model_context": list(model_context if model_context is not None else history),
        "checkpoint": checkpoint,
        "events": [],
        "next_seq": 1,
        "cancel_requested": False,
        "cancel_event": threading.Event(),
        "created_at": now,
        "updated_at": now,
        "started_at": 0.0,
        "finished_at": 0.0,
        "error": "",
        "condition": threading.Condition(),
        "thread": None,
    }
    with _runs_lock:
        _runs[run["id"]] = run
        _prune_runs_locked()
    return run


def _prune_runs_locked() -> None:
    if len(_runs) <= _RUN_RETENTION_LIMIT:
        return
    finished = sorted(
        [
            run
            for run in _runs.values()
            if str(run.get("status")) in _RUN_TERMINAL_STATES
        ],
        key=lambda run: float(run.get("updated_at") or 0),
    )
    for finished_run in finished[: max(0, len(_runs) - _RUN_RETENTION_LIMIT)]:
        _runs.pop(str(finished_run.get("id")), None)


def _max_active_runs() -> int:
    raw = os.environ.get("METIS_MAX_ACTIVE_RUNS", str(_RUN_DEFAULT_MAX_ACTIVE))
    try:
        return max(1, int(raw))
    except (TypeError, ValueError):
        return _RUN_DEFAULT_MAX_ACTIVE


def _active_run_count_locked() -> int:
    return sum(1 for run in _runs.values() if str(run.get("status")) in _RUN_ACTIVE_STATES)


def _get_run(run_id: str) -> Optional[Dict[str, Any]]:
    with _runs_lock:
        return _runs.get(str(run_id or ""))


def _run_public_payload(run: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "run_id": run.get("id", ""),
        "id": run.get("id", ""),
        "session_id": run.get("session_id", ""),
        "assistant_id": run.get("assistant_id", ""),
        "status": run.get("status", "unknown"),
        "phase": run.get("phase", ""),
        "cancel_requested": bool(run.get("cancel_requested")),
        "created_at": run.get("created_at", 0),
        "updated_at": run.get("updated_at", 0),
        "started_at": run.get("started_at", 0),
        "finished_at": run.get("finished_at", 0),
        "event_count": len(run.get("events") or []),
        "last_seq": int(run.get("next_seq") or 1) - 1,
        "error": run.get("error", ""),
    }


def _latest_session_run(session_id: str, *, active_only: bool = True) -> Optional[Dict[str, Any]]:
    with _runs_lock:
        candidates = [
            run
            for run in _runs.values()
            if run.get("session_id") == session_id
            and (not active_only or str(run.get("status")) in _RUN_ACTIVE_STATES)
        ]
    if not candidates:
        return None
    return max(candidates, key=lambda run: float(run.get("created_at") or 0))


def _set_run_status(run: Dict[str, Any], status: str, *, phase: str = "", error: str = "") -> None:
    with run["condition"]:
        run["status"] = status
        if phase:
            run["phase"] = phase
        if error:
            run["error"] = error
        run["updated_at"] = time.time()
        if status in _RUN_TERMINAL_STATES:
            run["finished_at"] = run["updated_at"]
        run["condition"].notify_all()


def _append_run_event(run: Dict[str, Any], payload: Dict[str, Any]) -> Dict[str, Any]:
    with run["condition"]:
        seq = int(run.get("next_seq") or 1)
        run["next_seq"] = seq + 1
        event = dict(payload)
        event["run_id"] = run["id"]
        event["runId"] = run["id"]
        event["session_id"] = run["session_id"]
        event["sessionId"] = run["session_id"]
        event["assistant_id"] = run["assistant_id"]
        event["assistantId"] = run["assistant_id"]
        event["seq"] = seq
        if isinstance(event.get("payload"), dict):
            event["payload"] = {
                **event["payload"],
                "run_id": run["id"],
                "session_id": run["session_id"],
                "assistant_id": run["assistant_id"],
                "seq": seq,
            }
        run["events"].append(event)
        run["updated_at"] = time.time()
        if event.get("kind") == "runtime_status":
            phase = str(event.get("phase") or event.get("payload", {}).get("phase") or "")
            if phase:
                run["phase"] = phase
        if event.get("kind") == "error":
            run["error"] = str(event.get("message") or event.get("payload", {}).get("message") or "")
        run["condition"].notify_all()
        return event


def _run_cancel_event() -> Dict[str, Any]:
    return _event_payload(
        ErrorEvent(
            code="RUN_CANCELLED",
            title="运行已取消",
            message="本次后台运行已取消。",
            hint="可以重新发送，或从会话历史继续。",
            recoverable=False,
        )
    )


def _sse_payloads_from_chunk(chunk: Any) -> List[Any]:
    text = chunk.decode("utf-8") if isinstance(chunk, bytes) else str(chunk)
    payloads: List[Any] = []
    for packet in text.split("\n\n"):
        if not packet.strip():
            continue
        for raw_line in packet.splitlines():
            line = raw_line.strip()
            if not line.startswith("data: "):
                continue
            payload = line[len("data: ") :]
            if payload == "[DONE]":
                payloads.append("[DONE]")
                continue
            try:
                payloads.append(json.loads(payload))
            except json.JSONDecodeError:
                payloads.append({"type": "error", "message": payload})
    return payloads


def _extract_user_message() -> Any:
    if request.method == "GET":
        return str(request.args.get("prompt") or request.args.get("message") or "").strip()
    data = request.get_json(silent=True) or {}
    message = data.get("message", data.get("prompt", ""))
    if isinstance(message, list):
        return message
    return str(message or "").strip()


def _side_chat_messages_from_request() -> List[Dict[str, str]]:
    data = request.get_json(silent=True) or {}
    raw_messages = data.get("messages")
    messages: List[Dict[str, str]] = []

    if isinstance(raw_messages, list):
        for item in raw_messages[-_SIDE_CHAT_MAX_MESSAGES:]:
            if not isinstance(item, dict):
                continue
            role = str(item.get("role") or "").strip().lower()
            if role not in {"system", "user", "assistant"}:
                continue
            content = _message_text(item.get("content") or item.get("text") or "").strip()
            if not content:
                continue
            messages.append({"role": role, "content": content[:_SIDE_CHAT_MAX_MESSAGE_CHARS]})

    if not messages:
        message = _message_text(data.get("message") or data.get("prompt") or "").strip()
        if message:
            messages.append({"role": "user", "content": message[:_SIDE_CHAT_MAX_MESSAGE_CHARS]})

    if not any(message["role"] == "user" for message in messages):
        return []

    if not messages or messages[0]["role"] != "system":
        messages.insert(0, {"role": "system", "content": _SIDE_CHAT_SYSTEM_PROMPT})
    else:
        messages[0] = {
            "role": "system",
            "content": f"{_SIDE_CHAT_SYSTEM_PROMPT}\n\n{messages[0]['content']}"[:_SIDE_CHAT_MAX_MESSAGE_CHARS],
        }
    return messages


def _side_chat_model_from_request() -> str:
    data = request.get_json(silent=True) or {}
    model = str(data.get("model") or "").strip()
    if not model or len(model) > 120:
        return ""
    return model


def _sse(payload: Any) -> str:
    return sse_data(payload)


def _sse_comment(comment: str) -> str:
    return f": {comment}\n\n"


def _error_response_payload(info: ErrorInfo) -> Dict[str, Any]:
    return {
        "error": info.message,
        "code": info.code,
        "title": info.title,
        "message": info.message,
        "hint": info.hint,
        "recoverable": info.recoverable,
        "status": info.status,
        "details": (info.details or "")[:1000],
    }


def _done_usage_payload(done_event: Optional[DoneEvent]) -> Dict[str, int]:
    if done_event is None:
        return {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
    return {
        "prompt_tokens": int(done_event.prompt_tokens or 0),
        "completion_tokens": int(done_event.completion_tokens or 0),
        "total_tokens": int(done_event.total_tokens or 0),
    }


def _event_payload(event: Any) -> Dict[str, Any]:
    return agent_event_payload(event)


def _error_event_from_exception(exc: Exception, *, recoverable: bool = False) -> ErrorEvent:
    info = classify_llm_error(exc, recoverable=recoverable)
    return ErrorEvent(
        message=info.message,
        recoverable=info.recoverable,
        code=info.code,
        title=info.title,
        hint=info.hint,
        status=info.status,
        details=info.details,
    )


def _stream_exception_response(
    exc: Exception,
    context: str,
    *,
    recoverable: bool = False,
    session_id: Optional[str] = None,
    history: Optional[List[Dict[str, Any]]] = None,
    mode: str = "auto",
) -> Any:
    logger.error("%s: %s", sanitize_for_log(context), sanitize_for_log(exc))
    yield _sse(_event_payload(_error_event_from_exception(exc, recoverable=recoverable)))
    if session_id and history is not None:
        if not _save_session_history(session_id, history=history, mode=mode):
            yield _sse(_disk_full_event())
    else:
        _save_active_session()
    yield _sse("[DONE]")


def _stream_agent_response(
    messages: List[Dict[str, Any]],
    config: AgentConfig,
    *,
    session_id: Optional[str] = None,
    history: Optional[List[Dict[str, Any]]] = None,
    mode: str = "auto",
    cancel_event: Optional[threading.Event] = None,
    checkpoint: Optional[CheckpointRecorder] = None,
) -> Any:
    """Stream a normal agent response with shared permission and compact handling."""
    run_history = list(history if history is not None else messages)
    working_messages = list(messages)
    pending_compact_event: Optional[CompactEvent] = None
    estimated_prompt_tokens = _estimate_prompt_tokens(working_messages, config.system_prompt)
    stage = compaction_stage(estimated_prompt_tokens, config.llm_model)
    if stage == 1:
        working_messages = _mask_observations(working_messages, keep_recent=6)
    elif stage >= 2:
        pending_compact_event = _trigger_auto_compact(config, aggressive=(stage == 3))
        working_messages = _mask_observations(
            working_messages,
            keep_recent=4 if stage == 3 else 6,
            aggressive=(stage == 3),
        )
    try:
        registry = get_registry()
        tool_schemas = registry.get_all_schemas(format="openai")
        if config.enabled_tools:
            enabled = set(config.enabled_tools)
            tool_schemas = [
                schema
                for schema in tool_schemas
                if ((schema.get("function") or {}).get("name") or "") in enabled
            ]
        valid_tool_schemas = [
            schema
            for schema in tool_schemas
            if (schema.get("function") or {}).get("name")
            and (schema.get("function") or {}).get("description")
            and isinstance((schema.get("function") or {}).get("parameters"), dict)
        ]
        logger.info(
            "Agent stream start: tools=%s valid_schemas=%s cwd=%s workspace_root=%s backend=%s model=%s",
            registry.tool_count,
            len(valid_tool_schemas),
            os.getcwd(),
            config.workspace_root,
            config.llm_backend,
            config.llm_model,
        )
        messages = working_messages
        if cancel_event is None:
            gen = run_stream(messages, config, registry=registry)
        else:
            gen = run_stream(messages, config, registry=registry, cancel_event=cancel_event)
    except OperationCancelled:
        yield _sse(_run_cancel_event())
        _sync_active_runtime_history(session_id, run_history, mode)
        if not _save_session_history(session_id, history=run_history, mode=mode):
            yield _sse(_disk_full_event())
        if checkpoint is not None:
            checkpoint.finalize("canceled")
        yield _sse("[DONE]")
        return
    except Exception as exc:
        logger.error("Agent stream creation failed: %s", sanitize_for_log(exc))
        yield _sse(_event_payload(_error_event_from_exception(exc, recoverable=False)))
        _sync_active_runtime_history(session_id, run_history, mode)
        if not _save_session_history(session_id, history=run_history, mode=mode):
            yield _sse(_disk_full_event())
        if checkpoint is not None:
            checkpoint.finalize("failed")
        yield _sse("[DONE]")
        return

    send_value: Optional[bool] = None
    tool_names: List[str] = []
    pending_tool_calls: Dict[str, Dict[str, Any]] = {}
    try:
        if pending_compact_event is not None:
            yield _sse(_event_payload(pending_compact_event))
        while True:
            if cancel_event is not None and cancel_event.is_set():
                raise OperationCancelled("Run cancelled")
            try:
                if send_value is not None:
                    event = gen.send(send_value)
                    send_value = None
                else:
                    event = next(gen)
            except OperationCancelled:
                yield _sse(_run_cancel_event())
                break
            except StopIteration:
                break
            except Exception as exc:
                logger.error("Agent stream failed: %s", sanitize_for_log(exc))
                yield _sse(_event_payload(_error_event_from_exception(exc, recoverable=False)))
                break

            lock: Optional[threading.Event] = None
            if isinstance(event, PermissionRequestEvent):
                lock = threading.Event()
                with _perm_dict_lock:
                    _permission_locks[event.request_id] = lock
                    _permission_contexts[event.request_id] = {
                        "request_id": event.request_id,
                        "call_id": event.call_id,
                        "tool": event.tool_name,
                        "arguments": event.arguments,
                        "workspace_root": config.workspace_root,
                        "created_at": time.time(),
                    }

            if isinstance(event, ToolCallEvent):
                if checkpoint is not None:
                    checkpoint.capture_tool_call(event.tool_name, event.arguments or {})
                tool_names.append(event.tool_name)
                pending_tool_calls[event.call_id or event.tool_name] = {
                    "name": event.tool_name,
                    "arguments": event.arguments or {},
                }

            if isinstance(event, ToolResultEvent):
                # FABLEADV-16: persist tool call + result into the transcript so
                # the UI can rebuild tool cards after reload/compaction. Stored
                # as a transcript-only record (metis_kind=tool) that is filtered
                # out before building the model request (agent_loop).
                call = pending_tool_calls.pop(event.call_id or event.tool_name, {})
                run_history.append(
                    _new_message(
                        "assistant",
                        "",
                        metis_kind="tool",
                        metis_tool={
                            "call_id": event.call_id or "",
                            "name": event.tool_name or call.get("name", ""),
                            "arguments": call.get("arguments", {}),
                            "result": _truncate_tool_record(event.result),
                            "status": "error" if _tool_result_is_error(event.result) else "success",
                        },
                    )
                )

            if isinstance(event, ContentEvent):
                run_history.append(_new_message("assistant", event.text))

            yield _sse(_event_payload(event))

            if isinstance(event, ToolCallEvent) and event.tool_name in _SUBAGENT_TOOLS:
                yield _sse(
                    {
                        "type": "subagent_start",
                        "task_id": event.call_id,
                        "name": event.tool_name,
                        "progress": 0,
                        "status": "running",
                    }
                )
            elif isinstance(event, ToolResultEvent) and event.tool_name in _SUBAGENT_TOOLS:
                yield _sse(
                    {
                        "type": "subagent_done",
                        "task_id": event.call_id,
                        "name": event.tool_name,
                        "progress": 100,
                        "status": "done",
                        "result": event.result,
                    }
                )

            if (
                isinstance(event, DoneEvent)
                and session_id == _runtime_state.active_session_id
                and should_auto_compact(
                    event.prompt_tokens,
                    config.llm_model,
                )
            ):
                compact_event = _trigger_auto_compact(
                    config,
                    aggressive=compaction_stage(event.prompt_tokens, config.llm_model) == 3,
                )
                if compact_event is not None:
                    yield _sse(_event_payload(compact_event))

            if isinstance(event, DoneEvent):
                learning_event = _maybe_record_learning(event, tool_names, session_id=session_id, history=run_history)
                if learning_event is not None:
                    yield _sse(learning_event)

            if isinstance(event, PermissionRequestEvent) and lock is not None:
                deadline = time.time() + 300
                while not lock.is_set() and time.time() < deadline:
                    if cancel_event is not None and cancel_event.is_set():
                        with _perm_dict_lock:
                            _permission_locks.pop(event.request_id, None)
                            _permission_contexts.pop(event.request_id, None)
                            _permission_results.pop(event.request_id, None)
                        raise OperationCancelled("Run cancelled")
                    lock.wait(timeout=0.2)
                with _perm_dict_lock:
                    if event.request_id in _permission_results:
                        send_value = _permission_results.pop(event.request_id, False)
                        context = {}
                    else:
                        send_value = False
                        context = dict(_permission_contexts.get(event.request_id, {}))
                    _permission_locks.pop(event.request_id, None)
                    _permission_contexts.pop(event.request_id, None)
                    _permission_results.pop(event.request_id, None)
                if context:
                    _append_permission_audit(
                        {
                            **context,
                            "action": "deny",
                            "approved": False,
                            "remember": "",
                            "source": "permission_timeout",
                        },
                        workspace_root=config.workspace_root,
                    )
    except (GeneratorExit, Exception) as exc:
        try:
            close_fn = getattr(gen, "close", None)
            if callable(close_fn):
                close_fn()
        except Exception:
            pass
        finally:
            _sync_active_runtime_history(session_id, run_history, mode)
            _save_session_history(session_id, history=run_history, mode=mode)
            if checkpoint is not None:
                checkpoint.finalize("aborted" if isinstance(exc, GeneratorExit) else "failed")
        if isinstance(exc, GeneratorExit):
            return
        raise
    _sync_active_runtime_history(session_id, run_history, mode)
    if not _save_session_history(session_id, history=run_history, mode=mode):
        yield _sse(_disk_full_event())
    if checkpoint is not None:
        checkpoint.finalize("done")
    yield _sse("[DONE]")


def _run_registry_worker(run_id: str) -> None:
    run = _get_run(run_id)
    if run is None:
        return
    run["started_at"] = time.time()
    _set_run_status(run, "running", phase="starting")
    try:
        config = _load_config_for_workspace(str(run.get("workspace_root") or ""))
    except Exception as exc:
        logger.error("Agent run configuration failed: %s", sanitize_for_log(exc))
        _append_run_event(run, _event_payload(_error_event_from_exception(exc, recoverable=False)))
        _set_run_status(run, "failed", phase="failed", error=str(exc))
        checkpoint = run.get("checkpoint")
        if isinstance(checkpoint, CheckpointRecorder):
            checkpoint.finalize("failed")
        return

    generator = _stream_agent_response(
        list(run.get("model_context") or run["history"]),
        config,
        session_id=str(run["session_id"]),
        history=list(run["history"]),
        mode=str(run.get("mode") or "auto"),
        cancel_event=run.get("cancel_event"),
        checkpoint=run.get("checkpoint") if isinstance(run.get("checkpoint"), CheckpointRecorder) else None,
    )
    saw_done = False
    try:
        for chunk in generator:
            if run.get("cancel_requested"):
                _append_run_event(run, _run_cancel_event())
                _set_run_status(run, "canceled", phase="canceled")
                close = getattr(generator, "close", None)
                if callable(close):
                    close()
                return
            for payload in _sse_payloads_from_chunk(chunk):
                if payload == "[DONE]":
                    saw_done = True
                    continue
                if isinstance(payload, dict):
                    _append_run_event(run, payload)
                    if payload.get("kind") == "error":
                        run["error"] = str(payload.get("message") or payload.get("payload", {}).get("message") or "")
                    if payload.get("kind") == "done":
                        saw_done = True
            if run.get("cancel_requested"):
                _append_run_event(run, _run_cancel_event())
                _set_run_status(run, "canceled", phase="canceled")
                close = getattr(generator, "close", None)
                if callable(close):
                    close()
                return
    except OperationCancelled:
        _append_run_event(run, _run_cancel_event())
        _set_run_status(run, "canceled", phase="canceled")
        return
    except Exception as exc:
        logger.error("Agent run worker failed: %s", sanitize_for_log(exc))
        _append_run_event(run, _event_payload(_error_event_from_exception(exc, recoverable=False)))
        _set_run_status(run, "failed", phase="failed", error=str(exc))
        return

    if run.get("cancel_requested"):
        _append_run_event(run, _run_cancel_event())
        _set_run_status(run, "canceled", phase="canceled")
    elif run.get("error"):
        _set_run_status(run, "failed", phase="failed")
    else:
        _set_run_status(run, "done", phase="completed" if saw_done else str(run.get("phase") or "done"))


def _start_run_thread(run: Dict[str, Any]) -> None:
    thread = threading.Thread(
        target=_run_registry_worker,
        args=(str(run["id"]),),
        daemon=True,
        name=f"metis-run-{str(run['id'])[:8]}",
    )
    run["thread"] = thread
    thread.start()


def _run_events_response(run: Dict[str, Any], after_seq: int = 0) -> Response:
    def stream() -> Any:
        last_seq = max(0, int(after_seq or 0))
        while True:
            with run["condition"]:
                events = [
                    event
                    for event in list(run.get("events") or [])
                    if int(event.get("seq") or 0) > last_seq
                ]
                status = str(run.get("status") or "")
                if not events and status not in _RUN_TERMINAL_STATES:
                    run["condition"].wait(timeout=15)
                    events = [
                        event
                        for event in list(run.get("events") or [])
                        if int(event.get("seq") or 0) > last_seq
                    ]
                    status = str(run.get("status") or "")
                    if not events and status not in _RUN_TERMINAL_STATES:
                        yield _sse_comment("heartbeat")
                        continue

            for event in events:
                last_seq = max(last_seq, int(event.get("seq") or 0))
                yield _sse(event)

            with run["condition"]:
                status = str(run.get("status") or "")
                has_more = any(int(event.get("seq") or 0) > last_seq for event in run.get("events") or [])
            if status in _RUN_TERMINAL_STATES and not has_more:
                yield _sse("[DONE]")
                return

    return _sse_response(stream())


def _sse_response(generator: Any) -> Response:
    def close_aware_stream() -> Any:
        try:
            for chunk in generator:
                if isinstance(chunk, bytes):
                    yield chunk
                else:
                    yield str(chunk).encode("utf-8")
        finally:
            close = getattr(generator, "close", None)
            if callable(close):
                close()

    return Response(
        stream_with_context(close_aware_stream()),
        content_type="text/event-stream; charset=utf-8",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        direct_passthrough=True,
    )


@app.errorhandler(413)
def request_entity_too_large(_: Any) -> Any:
    return jsonify({"error": "File too large. Max 10MB."}), 413


@app.route("/", methods=["GET"])
def index() -> Any:
    if os.path.exists(_FRONTEND_INDEX):
        return send_file(_FRONTEND_INDEX, mimetype="text/html; charset=utf-8")
    return (
        "<h1>Metis Agent</h1>"
        '<p>POST /chat or /chat/sync with {"message": "..."} to interact.</p>'
    )


@app.route("/health", methods=["GET"])
def health() -> Any:
    return jsonify({"ok": True})


@app.route("/contract/agent-events", methods=["GET"])
def agent_events_contract() -> Any:
    """Return the stream event contract consumed by the desktop app."""
    return jsonify(agent_event_contract_payload())


@app.route("/runs", methods=["GET"])
def list_runs() -> Any:
    session_id = str(request.args.get("session_id") or request.args.get("sessionId") or "").strip()
    with _runs_lock:
        runs = list(_runs.values())
    if session_id:
        runs = [run for run in runs if run.get("session_id") == session_id]
    runs.sort(key=lambda run: float(run.get("created_at") or 0), reverse=True)
    return jsonify({"runs": [_run_public_payload(run) for run in runs[:50]]})


@app.route("/runs", methods=["POST"])
def create_run() -> Any:
    user_message = _extract_user_message()
    if not user_message:
        return jsonify({"error": "No message provided"}), 400

    data = request.get_json(silent=True) or {}
    assistant_id = str(data.get("assistant_id") or data.get("assistantId") or "").strip() or f"assistant-run-{uuid.uuid4().hex[:12]}"
    session_id, history, compact_state, mode, explicit_session = _prepare_chat_session(user_message)
    active_for_session = _active_run_for_session(session_id)
    with _runs_lock:
        if active_for_session:
            return jsonify({"ok": False, "error": "session already has an active run", "run": _run_public_payload(active_for_session)}), 409
        if _active_run_count_locked() >= _max_active_runs():
            return jsonify({"ok": False, "error": "too many active runs", "max_active_runs": _max_active_runs()}), 429

    was_empty = not history
    workspace_root = _request_workspace_root(session_id)
    checkpoint_anchor = len(history)
    user_record = _new_message("user", user_message)
    history.append(user_record)
    checkpoint = create_checkpoint(
        session_id=session_id,
        workspace_root=workspace_root,
        anchor_index=checkpoint_anchor,
        user_message_id=str(user_record.get("id") or ""),
    )
    model_context = _model_context_with_skill_invocation(history, compact_state, workspace_root)
    _save_session_history(session_id, history=history, mode=mode)
    _commit_request_history_to_active(session_id, history, compact_state, mode, explicit_session)
    if was_empty and session_id:
        get_session_manager().update_session(session_id, title=_auto_title(_message_text(user_message)))
        _generate_smart_title(session_id, user_message)

    run = _create_run_state(
        session_id=session_id,
        assistant_id=assistant_id,
        history=history,
        mode=mode,
        model_context=model_context,
        workspace_root=workspace_root,
        checkpoint=checkpoint,
    )
    _start_run_thread(run)
    return jsonify({"ok": True, **_run_public_payload(run)})


@app.route("/runs/<run_id>", methods=["GET"])
def get_run(run_id: str) -> Any:
    run = _get_run(run_id)
    if run is None:
        return jsonify({"error": "run not found"}), 404
    return jsonify(_run_public_payload(run))


@app.route("/runs/<run_id>/events", methods=["GET"])
def run_events(run_id: str) -> Any:
    run = _get_run(run_id)
    if run is None:
        return jsonify({"error": "run not found"}), 404
    after_seq = int(request.args.get("after") or request.args.get("after_seq") or 0)
    return _run_events_response(run, after_seq)


@app.route("/runs/<run_id>/cancel", methods=["POST", "DELETE"])
def cancel_run(run_id: str) -> Any:
    run = _get_run(run_id)
    if run is None:
        return jsonify({"error": "run not found"}), 404
    should_emit = False
    with run["condition"]:
        if str(run.get("status")) not in _RUN_TERMINAL_STATES:
            run["cancel_requested"] = True
            cancel_event = run.get("cancel_event")
            if isinstance(cancel_event, threading.Event):
                cancel_event.set()
            run["status"] = "canceling"
            run["phase"] = "cancel_requested"
            run["updated_at"] = time.time()
            should_emit = True
            run["condition"].notify_all()
    if should_emit:
        _append_run_event(
            run,
            {
                "type": "runtime_status",
                "kind": "runtime_status",
                "payload": {
                    "phase": "cancel_requested",
                    "message": "Cancel requested",
                    "recoverable": False,
                },
            },
        )
    return jsonify({"ok": True, **_run_public_payload(run)})


@app.route("/sessions/<session_id>/checkpoints", methods=["GET"])
def session_checkpoints(session_id: str) -> Any:
    session = get_session_manager().get_session(session_id)
    if session is None:
        return jsonify({"error": "session not found"}), 404
    payload = []
    for item in reversed(list_checkpoints(session_id)):
        payload.append(
            {
                "checkpoint_id": item.get("checkpoint_id", ""),
                "session_id": item.get("session_id", session_id),
                "anchor_index": int(item.get("anchor_index") or 0),
                "user_message_id": item.get("user_message_id", ""),
                "reason": item.get("reason", "user_turn"),
                "created_at": item.get("created_at", 0),
                "completed_at": item.get("completed_at", 0),
                "status": item.get("status", ""),
                "file_count": len(item.get("files") or []),
                "files": [
                    {
                        "relative_path": file.get("relative_path", ""),
                        "existed": bool(file.get("existed")),
                        "skipped": file.get("skipped", ""),
                    }
                    for file in item.get("files", [])
                    if isinstance(file, dict)
                ],
            }
        )
    return jsonify({"checkpoints": payload})


@app.route("/sessions/<session_id>/rewind", methods=["POST"])
def session_rewind(session_id: str) -> Any:
    data = request.get_json(silent=True) or {}
    mode = str(data.get("mode") or "both").strip().lower()
    if mode not in {"conversation", "files", "both"}:
        return jsonify({"ok": False, "error": "mode must be conversation, files, or both"}), 400
    active_run = _active_run_for_session(session_id)
    if active_run is not None:
        return jsonify({"ok": False, "error": "当前会话仍有任务运行，结束后才能 rewind。", "run": _run_public_payload(active_run)}), 409

    manager = get_session_manager()
    session = manager.get_session(session_id)
    if session is None:
        return jsonify({"ok": False, "error": "session not found"}), 404

    checkpoint = find_checkpoint(
        session_id,
        checkpoint_id=str(data.get("checkpoint_id") or data.get("checkpointId") or ""),
        user_message_id=str(data.get("message_id") or data.get("messageId") or ""),
        anchor_index=data.get("anchor_index", data.get("anchorIndex")),
    )
    if checkpoint is None:
        return jsonify({"ok": False, "error": "没有可用 checkpoint。"}), 404

    workspace_root = _request_workspace_root(session_id)
    safety = create_checkpoint(
        session_id=session_id,
        workspace_root=workspace_root,
        anchor_index=len(session.history),
        reason="rewind_safety",
        history_snapshot=list(session.history),
        compact_state_snapshot=dict(session.compact_state),
        prune=False,
    )
    if mode in {"files", "both"}:
        for rel_path in _checkpoint_paths_from(checkpoint, session_id):
            safety.capture_path(rel_path)
    safety.finalize("rewind_safety", prune=False)

    restored: Dict[str, Any] = {"restored": [], "skipped": []}
    if mode in {"files", "both"}:
        restored = restore_files_from_checkpoint(session_id, str(checkpoint.get("checkpoint_id") or ""), workspace_root)

    next_history = list(session.history)
    next_compact_state = dict(session.compact_state)
    if mode in {"conversation", "both"}:
        snapshot_history = load_history_snapshot(session_id, checkpoint)
        snapshot_compact = load_compact_state_snapshot(session_id, checkpoint)
        if snapshot_history is not None:
            next_history = snapshot_history
            next_compact_state = snapshot_compact or {}
        else:
            anchor_index = max(0, min(int(checkpoint.get("anchor_index") or 0), len(session.history)))
            next_history = list(session.history[:anchor_index])
            next_compact_state = _compact_state_after_truncate(session.history, next_history, session.compact_state)
        manager.update_session(session_id, history=next_history, compact_state=next_compact_state, mode=session.mode)
        if session_id == _runtime_state.active_session_id:
            _runtime_state.chat_history = list(next_history)
            _runtime_state.compact_state = dict(next_compact_state)

    prune_checkpoints(session_id)

    return jsonify(
        {
            "ok": True,
            "mode": mode,
            "checkpoint_id": checkpoint.get("checkpoint_id", ""),
            "safety_checkpoint_id": safety.checkpoint_id,
            "history_length": len(next_history),
            "restored": restored.get("restored", []),
            "skipped": restored.get("skipped", []),
        }
    )


def _checkpoint_paths_from(checkpoint: Dict[str, Any], session_id: str) -> List[str]:
    checkpoints = list_checkpoints(session_id)
    selected_id = str(checkpoint.get("checkpoint_id") or "")
    start = next((index for index, item in enumerate(checkpoints) if item.get("checkpoint_id") == selected_id), -1)
    if start < 0:
        return []
    paths: List[str] = []
    for item in checkpoints[start:]:
        for file in item.get("files") or []:
            if not isinstance(file, dict):
                continue
            rel_path = str(file.get("relative_path") or "")
            if rel_path and rel_path not in paths:
                paths.append(rel_path)
    return paths


@app.route("/sessions/<session_id>/runs/active", methods=["GET"])
def active_session_run(session_id: str) -> Any:
    run = _latest_session_run(session_id, active_only=True)
    if run is None:
        return jsonify({"ok": False, "run": None})
    return jsonify({"ok": True, "run": _run_public_payload(run)})


@app.route("/side-chat", methods=["POST"])
def side_chat() -> Any:
    messages = _side_chat_messages_from_request()
    if not messages:
        return jsonify({"error": "No side chat message provided"}), 400

    try:
        config = build_agent_config(system_prompt="", execution_mode="auto")
        model_override = _side_chat_model_from_request()
        if model_override:
            config = replace(config, llm_model=model_override)
    except Exception as exc:
        logger.error("Side chat configuration failed: %s", sanitize_for_log(exc))
        error_payload = _event_payload(_error_event_from_exception(exc, recoverable=False))

        def config_error_stream() -> Any:
            yield _sse(error_payload)
            yield _sse("[DONE]")

        return _sse_response(config_error_stream())

    def stream() -> Any:
        cancel_event = threading.Event()
        llm_stream = None
        closed = False
        try:
            from backend.runtime.agent_loop import _create_backend

            backend = _create_backend(config)
            yield _sse(
                {
                    "schema": "metis.agent_event.v1",
                    "kind": "runtime_status",
                    "type": "runtime_status",
                    "payload": {
                        "phase": "llm_request",
                        "message": "Calling standalone side chat model",
                        "recoverable": True,
                    },
                }
            )
            llm_stream = backend.chat_stream(
                messages,
                tools=None,
                temperature=config.temperature,
                max_tokens=config.max_tokens,
                timeout=config.timeout,
                cancel_event=cancel_event,
            )
            yield _sse(
                {
                    "schema": "metis.agent_event.v1",
                    "kind": "runtime_status",
                    "type": "runtime_status",
                    "payload": {
                        "phase": "streaming",
                        "message": "Receiving standalone side chat response",
                        "recoverable": True,
                    },
                }
            )
            while True:
                try:
                    chunk = next(llm_stream)
                except StopIteration as stop:
                    response = stop.value
                    break
                if chunk:
                    yield _sse(_event_payload(TextDeltaEvent(text=chunk)))

            if response is not None and getattr(response, "thinking", ""):
                yield _sse(_event_payload(ThinkingEvent(text=str(response.thinking))))
            usage = getattr(response, "usage", None)
            yield _sse(
                _event_payload(
                    DoneEvent(
                        total_turns=1,
                        total_tool_calls=0,
                        prompt_tokens=int(getattr(usage, "prompt_tokens", 0) or 0),
                        completion_tokens=int(getattr(usage, "completion_tokens", 0) or 0),
                        total_tokens=int(getattr(usage, "total_tokens", 0) or 0),
                        prompt_cache_hit_tokens=int(getattr(usage, "prompt_cache_hit_tokens", 0) or 0),
                        prompt_cache_miss_tokens=int(getattr(usage, "prompt_cache_miss_tokens", 0) or 0),
                    )
                )
            )
        except GeneratorExit:
            closed = True
            cancel_event.set()
            close = getattr(llm_stream, "close", None)
            if callable(close):
                close()
            raise
        except Exception as exc:
            logger.error("Side chat stream failed: %s", sanitize_for_log(exc))
            yield _sse(_event_payload(_error_event_from_exception(exc, recoverable=False)))
        finally:
            if not closed:
                yield _sse("[DONE]")

    return _sse_response(stream())


@app.route("/chat", methods=["POST", "GET"])
def chat() -> Any:
    user_message = _extract_user_message()
    if not user_message:
        return jsonify({"error": "No message provided"}), 400

    session_id, history, compact_state, mode, explicit_session = _prepare_chat_session(user_message)
    workspace_root = _request_workspace_root(session_id)
    was_empty = not history
    checkpoint_anchor = len(history)
    user_record = _new_message("user", user_message)
    history.append(user_record)
    checkpoint = create_checkpoint(
        session_id=session_id,
        workspace_root=workspace_root,
        anchor_index=checkpoint_anchor,
        user_message_id=str(user_record.get("id") or ""),
    )
    _commit_request_history_to_active(session_id, history, compact_state, mode, explicit_session)
    if was_empty and session_id:
        get_session_manager().update_session(session_id, title=_auto_title(_message_text(user_message)))
        _generate_smart_title(session_id, user_message)

    try:
        config = _load_config_for_workspace(workspace_root)
    except Exception as exc:
        checkpoint.finalize("failed")
        return _sse_response(
            _stream_exception_response(
                exc,
                "Agent configuration failed",
                session_id=session_id,
                history=history,
                mode=mode,
            )
        )

    model_context = _model_context_with_skill_invocation(history, compact_state, workspace_root)
    return _sse_response(_stream_agent_response(model_context, config, session_id=session_id, history=history, mode=mode, checkpoint=checkpoint))


@app.route("/chat/edit", methods=["POST"])
def chat_edit() -> Any:
    """Edit a user message at an index and regenerate from that point."""
    data = request.get_json(silent=True) or {}
    index = data.get("index")
    new_message = data.get("message", "")

    if not isinstance(index, int):
        return jsonify({"error": "index required (integer)"}), 400
    if new_message in (None, ""):
        return jsonify({"error": "message required"}), 400
    if index < 0 or index >= len(_runtime_state.chat_history):
        return jsonify({"error": "index out of range"}), 400
    if _runtime_state.chat_history[index].get("role") != "user":
        return jsonify({"error": "can only edit user messages"}), 400

    original_history = list(_runtime_state.chat_history)
    _runtime_state.chat_history = _runtime_state.chat_history[:index]
    _runtime_state.chat_history.append(_new_message("user", new_message))
    _runtime_state.compact_state = _compact_state_after_truncate(
        original_history,
        _runtime_state.chat_history,
        _runtime_state.compact_state,
    )
    _save_active_session()
    try:
        config = _load_config_for_workspace(_active_workspace_root())
    except Exception as exc:
        return _sse_response(_stream_exception_response(exc, "Agent configuration failed"))
    model_context = _model_context_for_history(_runtime_state.chat_history, _runtime_state.compact_state)
    return _sse_response(
        _stream_agent_response(
            model_context,
            config,
            session_id=_runtime_state.active_session_id,
            history=list(_runtime_state.chat_history),
            mode=_runtime_state.execution_mode,
        )
    )


@app.route("/chat/regenerate", methods=["POST"])
def chat_regenerate() -> Any:
    """Remove the last assistant/tool response and regenerate from the last user message."""
    if not _runtime_state.chat_history:
        return jsonify({"error": "empty history"}), 400

    original_history = list(_runtime_state.chat_history)
    while _runtime_state.chat_history and _runtime_state.chat_history[-1].get("role") != "user":
        _runtime_state.chat_history.pop()
    if not _runtime_state.chat_history:
        return jsonify({"error": "no user message to regenerate from"}), 400

    _runtime_state.compact_state = _compact_state_after_truncate(
        original_history,
        _runtime_state.chat_history,
        _runtime_state.compact_state,
    )
    _save_active_session()
    try:
        config = _load_config_for_workspace(_active_workspace_root())
    except Exception as exc:
        return _sse_response(_stream_exception_response(exc, "Agent configuration failed"))
    model_context = _model_context_for_history(_runtime_state.chat_history, _runtime_state.compact_state)
    return _sse_response(
        _stream_agent_response(
            model_context,
            config,
            session_id=_runtime_state.active_session_id,
            history=list(_runtime_state.chat_history),
            mode=_runtime_state.execution_mode,
        )
    )


@app.route("/chat/sync", methods=["POST"])
def chat_sync() -> Any:
    user_message = _extract_user_message()
    if not user_message:
        return jsonify({"error": "No message provided"}), 400

    session_id, history, compact_state, mode, explicit_session = _prepare_chat_session(user_message)
    workspace_root = _request_workspace_root(session_id)
    was_empty = not history
    checkpoint_anchor = len(history)
    user_record = _new_message("user", user_message)
    history.append(user_record)
    checkpoint = create_checkpoint(
        session_id=session_id,
        workspace_root=workspace_root,
        anchor_index=checkpoint_anchor,
        user_message_id=str(user_record.get("id") or ""),
    )
    _commit_request_history_to_active(session_id, history, compact_state, mode, explicit_session)
    if was_empty and session_id:
        get_session_manager().update_session(session_id, title=_auto_title(_message_text(user_message)))
        _generate_smart_title(session_id, user_message)
    try:
        config = _load_config_for_workspace(workspace_root)
    except Exception as exc:
        logger.error("Agent configuration failed: %s", sanitize_for_log(exc))
        info = classify_llm_error(exc, recoverable=False)
        _save_session_history(session_id, history=history, mode=mode)
        checkpoint.finalize("failed")
        return jsonify(
            {
                "response": "",
                "tool_calls": [],
                "errors": [_error_response_payload(info)],
            }
        )
    messages = _model_context_with_skill_invocation(history, compact_state, workspace_root)
    final_text = ""
    tool_calls: List[Dict[str, Any]] = []
    errors: List[Dict[str, Any]] = []
    last_done: Optional[DoneEvent] = None

    for event in run(messages, config):
        if isinstance(event, ContentEvent):
            final_text = event.text
        elif isinstance(event, ToolCallEvent):
            checkpoint.capture_tool_call(event.tool_name, event.arguments or {})
            tool_calls.append(
                {"tool": event.tool_name, "args": event.arguments, "call_id": event.call_id}
            )
        elif isinstance(event, ErrorEvent):
            errors.append(
                {
                    "code": event.code,
                    "title": event.title,
                    "message": event.message,
                    "hint": event.hint,
                    "recoverable": event.recoverable,
                    "status": event.status,
                    "details": event.details,
                }
            )
        elif isinstance(event, DoneEvent):
            last_done = event

    if final_text:
        history.append(_new_message("assistant", final_text))
        _commit_request_history_to_active(session_id, history, compact_state, mode, explicit_session)
    _save_session_history(session_id, history=history, mode=mode)
    checkpoint.finalize("done" if not errors else "failed")
    if last_done and session_id == _runtime_state.active_session_id and should_auto_compact(last_done.prompt_tokens, config.llm_model):
        _trigger_auto_compact(config)

    return jsonify(
        {
            "response": final_text,
            "tool_calls": tool_calls,
            "errors": errors,
            "history_length": len(history),
            "usage": _done_usage_payload(last_done),
            "session_id": session_id,
        }
    )


@app.route("/status", methods=["GET"])
def status() -> Any:
    registry = get_registry()
    config = _load_config_for_workspace(_active_workspace_root())
    vision = False
    try:
        from backend.runtime.agent_loop import _create_backend

        vision = _create_backend(config).supports_vision
    except Exception:
        vision = False
    return jsonify(
        {
            "status": "running",
            "llm_backend": config.llm_backend,
            "llm_model": config.llm_model,
            "tools_count": registry.tool_count,
            "history_length": len(_runtime_state.chat_history),
            "vision": vision,
            "session_id": _runtime_state.active_session_id,
        }
    )


def _model_capabilities_payload(overrides: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
    from backend.bridges.model_capability import detect_from_model_name
    from backend.runtime.agent_loop import _create_backend
    from backend.runtime.tool_tiers import tools_for_tier

    registry = get_registry()
    config = _load_config_for_workspace(_active_workspace_root())
    overrides = overrides or {}
    if overrides:
        config = replace(
            config,
            llm_backend=overrides.get("backend") or config.llm_backend,
            llm_base_url=overrides.get("base_url") or config.llm_base_url,
            llm_model=overrides.get("model") or config.llm_model,
        )
    caps = detect_from_model_name(config.llm_model)
    forced_tier = os.environ.get("METIS_TOOL_TIER", "").strip()
    tier = caps.tier
    if forced_tier:
        try:
            tier = int(forced_tier)
        except ValueError:
            tier = caps.tier
    allowed = tools_for_tier(tier)
    supports_vision = False
    try:
        supports_vision = bool(_create_backend(config).supports_vision)
    except Exception:
        supports_vision = False
    total_tool_count = registry.tool_count
    return {
        "tier": tier,
        "tier_label": {1: "强", 2: "中", 3: "基础"}.get(tier, "中"),
        "tierLabel": {1: "强", 2: "中", 3: "基础"}.get(tier, "中"),
        "family": caps.detected_family,
        "model": caps.detected_model,
        "detection_method": caps.detection_method,
        "detectionMethod": caps.detection_method,
        "effective_context": caps.effective_context,
        "effectiveContext": caps.effective_context,
        "supports_vision": supports_vision,
        "supportsVision": supports_vision,
        "vision_protocol": caps.vision_protocol,
        "visionProtocol": caps.vision_protocol,
        "supports_tool_calling": caps.supports_tool_calling,
        "supportsToolCalling": caps.supports_tool_calling,
        "supports_structured_output": caps.supports_structured_output,
        "supportsStructuredOutput": caps.supports_structured_output,
        "instruction_adherence": caps.instruction_adherence,
        "instructionAdherence": caps.instruction_adherence,
        "tool_count": len(allowed) if allowed is not None else total_tool_count,
        "toolCount": len(allowed) if allowed is not None else total_tool_count,
        "total_tool_count": total_tool_count,
        "totalToolCount": total_tool_count,
    }


@app.route("/api/model/capabilities", methods=["GET"])
def model_capabilities() -> Any:
    return jsonify(
        _model_capabilities_payload(
            {
                "backend": request.args.get("backend", ""),
                "base_url": request.args.get("base_url", ""),
                "model": request.args.get("model", ""),
            }
        )
    )


@app.route("/runtime/state", methods=["GET"])
def runtime_state() -> Any:
    """Return the current desktop runtime selection without exposing message text."""
    return jsonify(_runtime_state.snapshot())




def _read_text(path: str) -> str:
    if not os.path.isfile(path):
        return ""
    try:
        with open(path, "r", encoding="utf-8") as handle:
            return handle.read()
    except OSError:
        return ""


def _write_text(path: str, content: str) -> None:
    directory = os.path.dirname(path)
    if directory:
        os.makedirs(directory, exist_ok=True)
    with open(path, "w", encoding="utf-8") as handle:
        handle.write(content)


def _append_project_memory(entry: str) -> str:
    paths = _helpers.memory_paths_payload()
    project_path = paths["project_path"]
    current = _read_text(project_path).rstrip()
    heading = "## Metis Learned Notes"
    if heading not in current:
        current = f"{current}\n\n{heading}".strip()
    _write_text(project_path, f"{current}\n{entry}\n")
    return project_path



def _create_skill_from_session(summary: str, tool_names: List[str]) -> Optional[str]:
    slug = _helpers.safe_skill_slug(summary)
    dest = _helpers.unique_skill_dir(slug)
    os.makedirs(dest, exist_ok=True)
    title = summary[:60].strip() or "Metis workflow"
    summary_inline = " ".join(str(summary or "").split())
    safe_summary = summary_inline.replace('"', "'")
    tool_line = ", ".join(sorted(set(tool_names))) if tool_names else "standard chat workflow"
    content = (
        "---\n"
        f"name: {_helpers.safe_skill_slug(title)}\n"
        f"description: \"复用已成功完成的 Metis 工作流：{safe_summary[:180]}\"\n"
        f"when_to_use: \"当未来任务类似：{safe_summary[:220]}\"\n"
        "---\n"
        f"# {title}\n\n"
        "## Trigger\n"
        f"Use this skill when a future task resembles: {summary}\n\n"
        "## Workflow\n"
        "1. Confirm the user's concrete goal and workspace context.\n"
        "2. Reuse the same project conventions and tools that worked previously.\n"
        "3. Verify the change with the narrowest reliable check before reporting back.\n\n"
        "## Prior Successful Tools\n"
        f"{tool_line}\n"
    )
    skill_path = os.path.join(dest, "SKILL.md")
    _write_text(skill_path, content)
    return skill_path


def _maybe_record_learning(
    done_event: DoneEvent,
    tool_names: List[str],
    *,
    session_id: Optional[str] = None,
    history: Optional[List[Dict[str, Any]]] = None,
) -> Optional[Dict[str, Any]]:
    target_session_id = session_id or _runtime_state.active_session_id
    target_history = history if history is not None else _runtime_state.chat_history
    if not target_session_id or target_session_id in _runtime_state.learning_nudged_sessions:
        return None
    tool_count = done_event.total_tool_calls or len(tool_names)
    meaningful_tool_names = [
        name for name in tool_names if name and name not in {"todo_write"}
    ]
    if tool_count <= 0 or not meaningful_tool_names:
        return None
    complex_enough = tool_count >= 2 or done_event.total_turns >= 3 or len(target_history) >= 6
    if not complex_enough:
        return None

    first_user = ""
    for message in target_history:
        if message.get("role") == "user":
            first_user = _message_text(message.get("content", ""))
            break
    summary = (first_user or "Metis workflow").replace("\n", " ").strip()
    if len(summary) > 120:
        summary = summary[:117] + "..."

    memory_count = 0
    skill_count = 0
    memory_path = ""
    skill_path = ""
    if env_bool("METIS_AUTO_MEMORY", "MIRO_AUTO_MEMORY", True):
        day = time.strftime("%Y-%m-%d %H:%M")
        tools = ", ".join(sorted(set(meaningful_tool_names))) if meaningful_tool_names else "none"
        memory_path = _append_project_memory(f"- {day}: {summary} (tools: {tools})")
        memory_count = 1

    skill_tool_count = len(meaningful_tool_names) if tool_names else tool_count
    if env_bool("METIS_AUTO_SKILLS", "MIRO_AUTO_SKILLS", True) and skill_tool_count >= 2:
        skill_path = _create_skill_from_session(summary, meaningful_tool_names) or ""
        skill_count = 1 if skill_path else 0

    if memory_count or skill_count:
        _runtime_state.learning_nudged_sessions.add(target_session_id)
        return {
            "type": "memory_nudge",
            "memory_count": memory_count,
            "skill_count": skill_count,
            "memory_path": memory_path,
            "skill_path": skill_path,
            "message": f"已沉淀 {memory_count} 条记忆"
            + (f"，生成 {skill_count} 个技能" if skill_count else ""),
        }
    return None


def _run_cron_task(task: Dict[str, Any]) -> Dict[str, Any]:
    cwd_before = os.getcwd()
    workspace_id = str(task.get("workspace_id") or "")
    workspace = get_workspace_manager().get_workspace(workspace_id) if workspace_id else None
    try:
        if workspace and os.path.isdir(workspace.path):
            os.chdir(workspace.path)
        prompt = str(task.get("prompt") or "").strip()
        if not prompt:
            return {"ok": False, "error": "empty prompt"}
        title = f"[Cron] {str(task.get('name') or 'Scheduled task')[:48]}"
        session = get_session_manager().create_session(title=title, workspace_id=workspace_id)
        history: List[Dict[str, Any]] = [_new_message("user", prompt)]
        root = workspace.path if workspace else ""
        config = build_agent_config(
            system_prompt=_load_system_prompt(),
            user_memory_text=_load_metis_md(root),
            execution_mode="auto",
            permission_checker=_permission_checker_for_workspace(root),
            tool_boundary_overrides=_tool_boundary_overrides_for_workspace(root),
            workspace_root=root,
        )
        final_text = ""
        errors: List[str] = []
        for event in run(history, config):
            if isinstance(event, ContentEvent):
                final_text = event.text
            elif isinstance(event, ErrorEvent):
                errors.append(event.message)
        if final_text:
            history.append(_new_message("assistant", final_text))
        elif errors:
            history.append(_new_message("assistant", "\n".join(errors)))
        get_session_manager().update_session(session.id, history=history, mode="auto")
        return {"ok": True, "session_id": session.id}
    finally:
        os.chdir(cwd_before)



# --- Session routes moved to session_routes.py ---



# --- Export routes moved to session_routes.py ---



# --- Session switch/delete/rename routes moved to session_routes.py ---


@app.route("/mode", methods=["GET"])
def get_mode() -> Any:
    return jsonify({"mode": _runtime_state.execution_mode})


@app.route("/mode", methods=["POST"])
def set_mode() -> Any:
    data = request.get_json(silent=True) or {}
    mode = data.get("mode", "auto")
    valid_modes = ("ask", "edit", "plan", "auto", "bypass")
    if mode not in valid_modes:
        return jsonify({"error": f"invalid mode, must be one of: {valid_modes}"}), 400
    _runtime_state.execution_mode = mode
    _save_active_session()
    return jsonify({"mode": _runtime_state.execution_mode})


@app.route("/permission", methods=["POST"])
def handle_permission() -> Any:
    """Approve or deny a pending tool execution."""
    data = request.get_json(silent=True) or {}
    request_id = str(data.get("request_id", ""))
    approved = bool(data.get("approved", False))
    remember = str(data.get("remember") or "").strip().lower()
    if remember not in {"", "allow", "deny"}:
        return jsonify({"error": "remember must be allow, deny, or empty"}), 400

    with _perm_dict_lock:
        lock = _permission_locks.get(request_id)
        context = dict(_permission_contexts.get(request_id, {}))
    if not lock:
        return jsonify({"error": "no pending permission request"}), 404

    tool_name = str(context.get("tool") or data.get("tool") or "")
    arguments = context.get("arguments")
    if not isinstance(arguments, dict):
        raw_args = data.get("args", data.get("arguments", {}))
        arguments = raw_args if isinstance(raw_args, dict) else {}
    call_id = str(context.get("call_id") or data.get("call_id") or "")
    workspace_root = str(context.get("workspace_root") or "")
    rule_id = ""
    if remember in {"allow", "deny"} and tool_name:
        rule = _create_permission_rule(
            tool=tool_name,
            action=remember,
            source="permission_dialog",
            workspace_root=workspace_root,
        )
        rule_id = str(rule.get("id") or "")

    _append_permission_audit(
        {
            "request_id": request_id,
            "call_id": call_id,
            "tool": tool_name,
            "arguments": arguments,
            "action": "allow" if approved else "deny",
            "approved": approved,
            "remember": remember,
            "rule_id": rule_id,
            "source": "permission_dialog",
        },
        workspace_root=workspace_root,
    )
    with _perm_dict_lock:
        current_lock = _permission_locks.get(request_id)
        if not current_lock:
            return jsonify({"error": "no pending permission request"}), 404
        _permission_results[request_id] = approved
        current_lock.set()
    return jsonify({"ok": True, "approved": approved, "remember": remember, "rule_id": rule_id})


@app.route("/permissions", methods=["GET"])
def list_permissions() -> Any:
    root = _active_workspace_root()
    paths = _permission_paths(root)
    rules = [_permission_rule_payload(rule) for rule in _permission_rules(root)]
    return jsonify(
        {
            "rules": rules,
            "audit": _read_permission_audit(limit=50, workspace_root=root),
            "path": paths["config"],
            "legacy_path": paths["legacy_config"],
            "audit_path": paths["audit"],
        }
    )


@app.route("/permissions", methods=["POST"])
def create_permission() -> Any:
    data = request.get_json(silent=True) or {}
    tool = str(data.get("tool") or "").strip()
    action = str(data.get("action") or "").strip().lower()
    args_match = data.get("args_match", {})
    if not tool:
        return jsonify({"error": "tool required"}), 400
    if action not in {"allow", "deny", "ask"}:
        return jsonify({"error": "action must be allow, deny, or ask"}), 400
    if not isinstance(args_match, dict):
        return jsonify({"error": "args_match must be an object"}), 400
    rule = _create_permission_rule(
        tool=tool,
        action=action,
        args_match=args_match,
        source=str(data.get("source") or "settings"),
        workspace_root=_active_workspace_root(),
    )
    return jsonify({"ok": True, "rule": _permission_rule_payload(rule)})


@app.route("/permissions/<rule_id>", methods=["DELETE"])
def delete_permission(rule_id: str) -> Any:
    root = _active_workspace_root()
    rules = _permission_rules(root)
    next_rules = [rule for rule in rules if str(rule.get("id") or "") != rule_id]
    if len(next_rules) == len(rules):
        return jsonify({"error": "rule not found"}), 404
    _write_permission_rules(next_rules, root)
    return jsonify({"ok": True, "id": rule_id})



# --- Workspace CRUD routes moved to workspace_routes.py ---



# --- Workspace switch/delete/open-folder routes moved to workspace_routes.py ---


# --- workspace/tree, git-status, file-changes/revert moved to workspace_routes.py ---


@app.route("/reset", methods=["POST"])
def reset() -> Any:
    from backend.core.memory.workspace_state import clear_read_tracking

    _save_active_session()
    session = get_session_manager().create_session(workspace_id=_runtime_state.active_workspace_id or "")
    _runtime_state.activate_session(
        session.id,
        history=[],
        compact_state=dict(session.compact_state),
        mode=session.mode,
    )
    clear_read_tracking()
    return jsonify({"status": "reset", "session_id": session.id, "workspace_id": session.workspace_id})


@app.route("/compact", methods=["POST"])
def compact() -> Any:
    """Compact model context while preserving the persisted user transcript."""
    session_id = _request_session_id() or _runtime_state.active_session_id or ""
    if not session_id:
        payload = _compact_status_payload(
            ok=False,
            before_context_messages=0,
            after_context_messages=0,
            error="No active session to compact",
        )
        _runtime_state.last_compact_status = payload
        return jsonify(payload), 400

    active_run = _active_run_for_session(session_id)
    if active_run is not None:
        payload = _compact_status_payload(
            ok=False,
            before_context_messages=0,
            after_context_messages=0,
            error="请等待当前任务完成后再压缩上下文。",
        )
        _runtime_state.last_compact_status = payload
        return jsonify({**payload, "run": _run_public_payload(active_run)}), 409

    session = get_session_manager().get_session(session_id)
    if session is None:
        payload = _compact_status_payload(
            ok=False,
            before_context_messages=0,
            after_context_messages=0,
            error="session not found",
        )
        _runtime_state.last_compact_status = payload
        return jsonify(payload), 404

    history = list(session.history)
    current_state = _normalize_compact_state(session.compact_state)
    before_context = _model_context_for_history(history, current_state)
    keep_recent = 4
    if len(before_context) <= keep_recent + 2:
        payload = _compact_status_payload(
            ok=False,
            before_context_messages=len(before_context),
            after_context_messages=len(before_context),
            error="History too short to compact",
        )
        _runtime_state.last_compact_status = payload
        return jsonify(payload)

    result = _compact_history(before_context, keep_recent=keep_recent)
    if result is None:
        result = _mechanical_compact(before_context, keep_recent=keep_recent)

    summary = _compact_summary_from_messages(result)
    boundary_index = max(0, len(history) - keep_recent)
    boundary_message_id = ""
    if boundary_index < len(history):
        boundary_message_id = str(history[boundary_index].get("id") or "")
    compact_state = {
        "summary": summary,
        "boundary_message_id": boundary_message_id,
        "boundary_index": boundary_index,
        "compacted_at": time.time(),
        "compact_count": _safe_int(current_state.get("compact_count"), 0) + 1,
    }
    get_session_manager().update_session(session_id, compact_state=compact_state)
    if session_id == _runtime_state.active_session_id:
        _runtime_state.chat_history = history
        _runtime_state.compact_state = dict(compact_state)

    after_context = _model_context_for_history(history, compact_state)
    payload = _compact_status_payload(
        ok=True,
        before_context_messages=len(before_context),
        after_context_messages=len(after_context),
        summary_preview=summary[:200],
    )
    payload.update(
        {
            "history_count": len(history),
            "boundary_index": boundary_index,
            "boundary_message_id": boundary_message_id,
            "compact_count": compact_state["compact_count"],
        }
    )
    _runtime_state.last_compact_status = payload
    return jsonify(payload)


@app.route("/compact/status", methods=["GET"])
def compact_status() -> Any:
    return jsonify(_runtime_state.last_compact_status or {"running": False})


@app.route("/tools", methods=["GET"])
def list_tools() -> Any:
    registry = get_registry()
    from backend.runtime.tool_registry import apply_user_tool_config

    tool_config = apply_user_tool_config(registry)
    tools = [
        {
            "name": tool.name,
            "description": tool.description[:200],
            "source": tool.source,
            "toolset": profile.toolset if profile else tool.toolset,
            "available": registry.is_available(tool.name),
            "requires_approval": tool.requires_approval,
            "destructive": tool.destructive,
        }
        for tool in registry._tools.values()
        if registry.is_available(tool.name)
        for profile in [registry.get_tool_profile(tool.name)]
    ]
    return jsonify({"tools": tools, "count": len(tools), "tool_config": tool_config})


@app.route("/mcp/status", methods=["GET"])
def mcp_status() -> Any:
    """Return MCP server status and config sources."""
    try:
        from backend.runtime.mcp_client import MCPManager, get_mcp_manager
    except ImportError:
        return jsonify({"available": False})

    manager = get_mcp_manager()
    if not manager:
        return jsonify({
            "available": True,
            "enabled": not env_disabled("METIS_DISABLE_MCP", "MIRO_DISABLE_MCP"),
            "servers": {},
            "config_sources": MCPManager().get_config_sources(),
        })

    return jsonify({
        "available": True,
        "enabled": not env_disabled("METIS_DISABLE_MCP", "MIRO_DISABLE_MCP"),
        "servers": manager.get_status(),
        "config_sources": manager.get_config_sources(),
    })


@app.route("/mcp/reload", methods=["POST"])
def mcp_reload() -> Any:
    """Reload MCP config and refresh MCP-backed tools without restarting."""
    from backend.runtime.tool_registry import reload_mcp_tools

    data = request.get_json(silent=True) or {}
    config_path = str(data.get("config_path") or data.get("configPath") or request.args.get("config_path") or "")
    result = reload_mcp_tools(config_path=config_path)
    status = 200 if result.get("ok") else 500
    return jsonify(result), status


@app.route("/mcp/reconnect", methods=["POST"])
def mcp_reconnect() -> Any:
    """Reconnect a specific MCP server."""
    from backend.runtime.mcp_client import get_mcp_manager

    data = request.get_json(silent=True) or {}
    server_name = data.get("server")
    if not server_name:
        return jsonify({"error": "missing server name"}), 400

    manager = get_mcp_manager()
    if not manager:
        return jsonify({"error": "MCP not initialized"}), 503

    return jsonify(manager.reconnect(server_name))


@app.route("/mcp/disconnect", methods=["POST"])
def mcp_disconnect() -> Any:
    """Disconnect a specific MCP server."""
    from backend.runtime.mcp_client import get_mcp_manager

    data = request.get_json(silent=True) or {}
    server_name = data.get("server")
    if not server_name:
        return jsonify({"error": "missing server name"}), 400

    manager = get_mcp_manager()
    if not manager:
        return jsonify({"error": "MCP not initialized"}), 503

    return jsonify(manager.disconnect_one(server_name))


@app.route("/mcp/resources", methods=["GET"])
def mcp_resources() -> Any:
    """List resources exposed by one or all connected MCP servers."""
    from backend.runtime.mcp_client import get_mcp_manager

    manager = get_mcp_manager()
    if not manager:
        get_registry()
        manager = get_mcp_manager()
    if not manager:
        return jsonify({"error": "MCP not initialized"}), 503

    server_name = str(request.args.get("server") or "")
    resources = manager.list_resources(server_name)
    return jsonify({"ok": True, "resources": resources})


@app.route("/mcp/resources/read", methods=["POST"])
def mcp_read_resource() -> Any:
    """Read a single MCP resource from a connected server."""
    from backend.runtime.mcp_client import get_mcp_manager

    data = request.get_json(silent=True) or {}
    server_name = str(data.get("server") or "")
    uri = str(data.get("uri") or "")
    if not server_name or not uri:
        return jsonify({"error": "missing server or uri"}), 400

    manager = get_mcp_manager()
    if not manager:
        get_registry()
        manager = get_mcp_manager()
    if not manager:
        return jsonify({"error": "MCP not initialized"}), 503
    return jsonify(manager.read_resource(server_name, uri))



@app.route("/upload/parse", methods=["POST"])
def upload_parse() -> Any:
    """Parse an uploaded rich file and return extracted text."""
    if not _request_client_is_loopback():
        return jsonify({"error": "forbidden"}), 403
    if "file" not in request.files:
        return jsonify({"error": "no file uploaded"}), 400

    file = request.files["file"]
    filename = file.filename or "unknown"
    ext = os.path.splitext(filename)[1].lower()

    max_size = 10 * 1024 * 1024
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > max_size:
        return jsonify({"error": f"File too large ({size // 1024}KB). Max 10MB."}), 400

    try:
        if ext == ".pdf":
            text = _parse_pdf(file)
        elif ext == ".docx":
            text = _parse_docx(file)
        elif ext in (".xlsx", ".xls"):
            text = _parse_xlsx(file)
        elif ext == ".csv":
            text = file.read().decode("utf-8", errors="replace")
        elif ext in {
            ".txt",
            ".md",
            ".markdown",
            ".log",
            ".json",
            ".yaml",
            ".yml",
            ".toml",
            ".xml",
            ".py",
            ".js",
            ".ts",
            ".html",
            ".css",
        }:
            text = file.read().decode("utf-8", errors="replace")
        else:
            return jsonify({"error": f"Unsupported file type: {ext}"}), 400
    except ImportError as exc:
        return jsonify(
            {
                "error": f"Missing dependency for {ext} parsing. Run: pip install {_dep_for_ext(ext)}",
                "detail": str(exc),
            }
        ), 503
    except Exception as exc:
        return jsonify({"error": f"Failed to parse {filename}: {exc}"}), 500

    max_chars = 50000
    original_count = len(text)
    truncated = original_count > max_chars
    if truncated:
        text = text[:max_chars] + f"\n\n[...truncated, {original_count} total characters...]"

    return jsonify(
        {
            "filename": filename,
            "type": ext,
            "text": text,
            "char_count": original_count,
            "truncated": truncated,
        }
    )


def _parse_pdf(file_obj: Any) -> str:
    try:
        import PyPDF2
    except ImportError:
        try:
            import pypdf as PyPDF2
        except ImportError as exc:
            raise ImportError("PyPDF2 or pypdf") from exc

    reader = PyPDF2.PdfReader(file_obj)
    pages = []
    for index, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"--- Page {index + 1} ---\n{page_text.strip()}")
    return "\n\n".join(pages) or "(No extractable text found in PDF)"


def _parse_docx(file_obj: Any) -> str:
    import docx

    document = docx.Document(file_obj)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts) or "(No text found in document)"


def _parse_xlsx(file_obj: Any) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(file_obj, read_only=True, data_only=True)
    try:
        sheets = []
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            rows = []
            for row in worksheet.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                if any(cell.strip() for cell in cells):
                    rows.append(",".join(cells))
            if rows:
                sheets.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
        return "\n\n".join(sheets) or "(No data found in spreadsheet)"
    finally:
        workbook.close()


def _dep_for_ext(ext: str) -> str:
    deps = {
        ".pdf": "PyPDF2",
        ".docx": "python-docx",
        ".xlsx": "openpyxl",
        ".xls": "openpyxl",
    }
    return deps.get(ext, "unknown")



# --- file-preview and workspace/file routes moved to workspace_routes.py ---

def _request_client_is_loopback() -> bool:
    from backend.web.helpers import request_client_is_loopback
    return request_client_is_loopback()


@app.route("/internal/miro/open-files", methods=["POST"])
@app.route("/internal/metis/open-files", methods=["POST"])
def internal_miro_open_files() -> Any:
    if not _request_client_is_loopback():
        return jsonify({"error": "forbidden", "detail": "only loopback"}), 403
    body = request.get_json(force=False, silent=True)
    if body is None:
        return jsonify({"error": "invalid_json", "detail": "expect application/json"}), 400

    try:
        from backend.core.engine.constants import CURRENT_WORKSPACE
        from backend.core.engine.open_files_http_cache import set_http_open_files
    except Exception as exc:
        return jsonify({"error": "unavailable", "detail": str(exc)}), 503

    workspace = os.path.abspath(CURRENT_WORKSPACE)
    set_http_open_files(workspace, body)
    return jsonify({"ok": True, "workspace": workspace})


from backend.web.settings_routes import settings_bp
from backend.web.feature_routes import feature_bp
from backend.web.session_routes import session_bp
from backend.web.workspace_routes import workspace_bp
app.register_blueprint(settings_bp)
app.register_blueprint(feature_bp)
app.register_blueprint(session_bp)
app.register_blueprint(workspace_bp)

try:
    from backend.web.desk_blueprint import desk_bp

    app.register_blueprint(desk_bp)
except Exception:
    try:
        from desk_blueprint import desk_bp

        app.register_blueprint(desk_bp)
    except Exception as exc:
        logger.warning("desk_blueprint unavailable: %s", sanitize_for_log(exc))


_init_session()

try:
    from backend.web.scheduler import start_scheduler

    start_scheduler(_run_cron_task)
except Exception as exc:
    logger.warning("scheduler unavailable: %s", sanitize_for_log(exc))


def start_server(preferred_port: int = 5000, *, host: str = "127.0.0.1", max_attempts: int = 10) -> None:
    for offset in range(max_attempts):
        candidate = preferred_port + offset
        try:
            app.run(host=host, port=candidate, threaded=True, debug=False)
            return
        except OSError as exc:
            message = str(exc)
            if "Address already in use" in message or "10048" in message:
                logger.warning("Port %s is in use, trying %s", candidate, candidate + 1)
                continue
            raise
    raise RuntimeError(f"No available port in range {preferred_port}-{preferred_port + max_attempts - 1}")


if __name__ == "__main__":
    host = env("METIS_HTTP_HOST", "MIRO_HTTP_HOST", "127.0.0.1")
    port = int(env_any(["METIS_HTTP_PORT", "METIS_PORT", "MIRO_HTTP_PORT", "MIRO_PORT"], "5000"))
    config = _load_config_for_workspace(_active_workspace_root())
    registry = get_registry()
    logger.info("Metis Agent starting on http://%s:%s", host, port)
    logger.info("LLM: %s / %s", config.llm_backend, config.llm_model)
    logger.info("Tools: %s registered", registry.tool_count)
    start_server(port, host=host)

