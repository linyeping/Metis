from __future__ import annotations

import csv
import html.parser
import io
import re
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, List
from xml.etree import ElementTree as ET

from backend.runtime.document_converters import (
    antiword_path,
    document_converter_status,
    soffice_path,
)


MAX_UPLOAD_PARSE_CHARS = 50_000


class UnsupportedFileType(ValueError):
    """Raised when Metis has no safe parser for a file type."""


class MissingParserDependency(ImportError):
    """Raised when a supported parser needs an optional dependency."""

    def __init__(self, ext: str, dependency: str, detail: str = "") -> None:
        self.ext = ext
        self.dependency = dependency
        self.detail = detail
        super().__init__(detail or dependency)


@dataclass(frozen=True)
class ExtractedFile:
    filename: str
    extension: str
    text: str


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".log",
    ".json",
    ".jsonl",
    ".yaml",
    ".yml",
    ".toml",
    ".xml",
    ".ini",
    ".cfg",
    ".conf",
    ".env",
    ".py",
    ".pyw",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".mjs",
    ".cjs",
    ".html",
    ".htm",
    ".css",
    ".scss",
    ".less",
    ".sql",
    ".sh",
    ".bash",
    ".zsh",
    ".ps1",
    ".bat",
    ".cmd",
    ".java",
    ".go",
    ".rs",
    ".c",
    ".cc",
    ".cpp",
    ".h",
    ".hpp",
    ".cs",
    ".php",
    ".rb",
    ".swift",
    ".kt",
    ".kts",
    ".dart",
    ".vue",
    ".svelte",
    ".csv",
    ".tsv",
    ".eml",
}

LEGACY_OFFICE_EXTENSIONS = {".doc", ".ppt", ".xls"}


def extract_uploaded_file(filename: str, data: bytes) -> ExtractedFile:
    name = filename or "unknown"
    ext = Path(name).suffix.lower()
    if not ext and _looks_like_text(data):
        ext = ".txt"

    if ext in TEXT_EXTENSIONS:
        text = _decode_text(data)
        if ext in {".html", ".htm"}:
            text = _html_to_text(text)
        elif ext in {".csv", ".tsv"}:
            text = _normalize_delimited_text(text, delimiter="\t" if ext == ".tsv" else ",")
        return ExtractedFile(name, ext, text)
    if ext == ".rtf":
        return ExtractedFile(name, ext, _rtf_to_text(_decode_text(data)))
    if ext == ".pdf":
        return ExtractedFile(name, ext, _extract_pdf(data))
    if ext == ".docx":
        return ExtractedFile(name, ext, _extract_docx(data))
    if ext in {".xlsx", ".xlsm", ".xltx"}:
        return ExtractedFile(name, ext, _extract_xlsx(data))
    if ext == ".pptx":
        return ExtractedFile(name, ext, _extract_pptx(data))
    if ext in {".odt", ".odp", ".ods"}:
        return ExtractedFile(name, ext, _extract_opendocument(data))
    if ext == ".epub":
        return ExtractedFile(name, ext, _extract_epub(data))
    if ext in LEGACY_OFFICE_EXTENSIONS:
        return ExtractedFile(name, ext, _extract_legacy_office(data, ext, name))

    raise UnsupportedFileType(
        f"Unsupported file type: {ext or '(unknown)'}. Supported: txt/md/html/json/csv/tsv, "
        "pdf, docx, xlsx/xlsm, pptx, odt/odp/ods, epub, rtf. Legacy doc/xls/ppt "
        "need LibreOffice or antiword-style converters installed."
    )


def dependency_for_extension(ext: str) -> str:
    deps = {
        ".pdf": "pypdf or PyPDF2",
        ".docx": "python-docx",
        ".xlsx": "openpyxl",
        ".xlsm": "openpyxl",
        ".xltx": "openpyxl",
        ".doc": "LibreOffice soffice or antiword",
        ".ppt": "LibreOffice soffice",
        ".xls": "xlrd or LibreOffice soffice",
    }
    return deps.get(ext.lower(), "built-in parser")


def truncate_text(text: str, max_chars: int = MAX_UPLOAD_PARSE_CHARS) -> tuple[str, int, bool]:
    original_count = len(text)
    if original_count <= max_chars:
        return text, original_count, False
    return text[:max_chars] + f"\n\n[...truncated, {original_count} total characters...]", original_count, True


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "utf-16-le", "utf-16-be", "gb18030", "cp1252"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _looks_like_text(data: bytes) -> bool:
    if not data:
        return True
    sample = data[:4096]
    if b"\x00" in sample:
        return False
    decoded = sample.decode("utf-8", errors="ignore")
    return bool(decoded.strip())


def _normalize_delimited_text(text: str, *, delimiter: str) -> str:
    try:
        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        rows = []
        for index, row in enumerate(reader):
            if index >= 2000:
                rows.append("[...table truncated after 2000 rows...]")
                break
            rows.append("\t".join(cell.strip() for cell in row))
        return "\n".join(rows)
    except csv.Error:
        return text


class _HTMLTextExtractor(html.parser.HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: List[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style", "noscript"}:
            self._skip_depth += 1
        if tag in {"p", "br", "div", "section", "article", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._skip_depth:
            self._skip_depth -= 1
        if tag in {"p", "div", "section", "article", "li", "tr"}:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        value = " ".join(data.split())
        if value:
            self._parts.append(value)
            self._parts.append(" ")

    def text(self) -> str:
        value = "".join(self._parts)
        lines = [" ".join(line.split()) for line in value.splitlines()]
        return "\n".join(line for line in lines if line).strip()


def _html_to_text(text: str) -> str:
    parser = _HTMLTextExtractor()
    try:
        parser.feed(text)
        result = parser.text()
        return result or text
    except Exception:
        return text


def _rtf_to_text(text: str) -> str:
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = text.replace(r"\par", "\n")
    text = text.replace("{", "").replace("}", "")
    return "\n".join(line.strip() for line in text.splitlines() if line.strip()) or text


def _extract_pdf(data: bytes) -> str:
    errors: List[str] = []
    try:
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = []
        for index, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"--- Page {index + 1} ---\n{page_text.strip()}")
        if pages:
            return "\n\n".join(pages)
    except ImportError as exc:
        raise MissingParserDependency(".pdf", "pypdf or PyPDF2", str(exc)) from exc
    except Exception as exc:
        errors.append(f"pypdf: {exc}")

    try:
        import PyPDF2

        reader = PyPDF2.PdfReader(io.BytesIO(data))
        pages = []
        for index, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"--- Page {index + 1} ---\n{page_text.strip()}")
        if pages:
            return "\n\n".join(pages)
    except ImportError:
        pass
    except Exception as exc:
        errors.append(f"PyPDF2: {exc}")

    if errors:
        raise ValueError("; ".join(errors))
    return "(No extractable text found in PDF)"


def _extract_docx(data: bytes) -> str:
    errors: List[str] = []
    try:
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))
        if parts:
            return "\n\n".join(parts)
    except ImportError as exc:
        raise MissingParserDependency(".docx", "python-docx", str(exc)) from exc
    except Exception as exc:
        errors.append(f"python-docx: {exc}")

    try:
        text = _extract_word_ooxml(data)
        if text.strip():
            return text
    except Exception as exc:
        errors.append(f"ooxml: {exc}")

    if errors:
        raise ValueError("; ".join(errors))
    return "(No text found in document)"


def _extract_word_ooxml(data: bytes) -> str:
    targets = (
        "word/document.xml",
        "word/footnotes.xml",
        "word/endnotes.xml",
        "word/comments.xml",
    )
    parts: List[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        names = set(archive.namelist())
        for name in targets:
            if name in names:
                parts.extend(_paragraphs_from_word_xml(archive.read(name)))
        for name in sorted(n for n in names if re.match(r"word/(header|footer)\d+\.xml$", n)):
            parts.extend(_paragraphs_from_word_xml(archive.read(name)))
    return "\n\n".join(_dedupe_nonempty(parts))


def _paragraphs_from_word_xml(xml_bytes: bytes) -> List[str]:
    root = ET.fromstring(xml_bytes)
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraphs: List[str] = []
    for paragraph in root.iter(f"{ns}p"):
        pieces: List[str] = []
        for node in paragraph.iter():
            if node.tag == f"{ns}t" and node.text:
                pieces.append(node.text)
            elif node.tag == f"{ns}tab":
                pieces.append("\t")
            elif node.tag in {f"{ns}br", f"{ns}cr"}:
                pieces.append("\n")
        value = "".join(pieces).strip()
        if value:
            paragraphs.append(value)
    return paragraphs


def _extract_xlsx(data: bytes) -> str:
    try:
        import openpyxl
    except ImportError as exc:
        raise MissingParserDependency(".xlsx", "openpyxl", str(exc)) from exc

    workbook = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        sheets = []
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            rows = []
            for index, row in enumerate(worksheet.iter_rows(values_only=True)):
                if index >= 2000:
                    rows.append("[...sheet truncated after 2000 rows...]")
                    break
                cells = [_stringify_cell(cell) for cell in row]
                if any(cell.strip() for cell in cells):
                    rows.append("\t".join(cells))
            if rows:
                sheets.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
        return "\n\n".join(sheets) or "(No data found in spreadsheet)"
    finally:
        workbook.close()


def _stringify_cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _extract_pptx(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        slide_names = sorted(
            (name for name in archive.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", name)),
            key=_slide_sort_key,
        )
        slides = []
        for index, name in enumerate(slide_names, start=1):
            texts = _texts_from_drawing_xml(archive.read(name))
            if texts:
                slides.append(f"--- Slide {index} ---\n" + "\n".join(texts))
        return "\n\n".join(slides) or "(No text found in presentation)"


def _slide_sort_key(name: str) -> tuple[int, str]:
    match = re.search(r"slide(\d+)\.xml$", name)
    return (int(match.group(1)) if match else 0, name)


def _texts_from_drawing_xml(xml_bytes: bytes) -> List[str]:
    root = ET.fromstring(xml_bytes)
    text_tag = "{http://schemas.openxmlformats.org/drawingml/2006/main}t"
    return _dedupe_nonempty(node.text or "" for node in root.iter(text_tag))


def _extract_opendocument(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        if "content.xml" not in archive.namelist():
            raise ValueError("OpenDocument content.xml is missing")
        root = ET.fromstring(archive.read("content.xml"))
        parts: List[str] = []
        for node in root.iter():
            if not node.text:
                continue
            if node.tag.endswith(("}p", "}h", "}span")):
                value = " ".join(node.itertext()).strip()
                if value:
                    parts.append(value)
        return "\n\n".join(_dedupe_nonempty(parts)) or "(No text found in OpenDocument file)"


def _extract_epub(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        html_names = sorted(
            name
            for name in archive.namelist()
            if name.lower().endswith((".html", ".xhtml", ".htm"))
        )
        parts = []
        for name in html_names[:80]:
            try:
                parts.append(f"--- {name} ---\n{_html_to_text(_decode_text(archive.read(name)))}")
            except Exception:
                continue
        return "\n\n".join(part for part in parts if part.strip()) or "(No text found in EPUB)"


def _extract_legacy_office(data: bytes, ext: str, filename: str) -> str:
    if ext == ".xls":
        try:
            return _extract_xls_with_xlrd(data)
        except MissingParserDependency:
            pass

    if ext == ".doc":
        antiword = antiword_path()
        if antiword:
            return _run_converter_stdout([antiword], data, ext, filename)

    converted = _convert_with_soffice(data, ext, filename)
    if converted:
        return converted

    status = document_converter_status()
    hints = "; ".join(status.to_dict().get("hints", []))
    raise UnsupportedFileType(
        f"Legacy Office {ext} parsing needs {dependency_for_extension(ext)}. "
        "Metis checked configured portable converters, PATH, and Python modules. "
        f"{hints or 'Please convert the file to docx/xlsx/pptx, or install a converter and retry.'}"
    )


def _extract_xls_with_xlrd(data: bytes) -> str:
    try:
        import xlrd  # type: ignore
    except ImportError as exc:
        raise MissingParserDependency(".xls", "xlrd", str(exc)) from exc

    workbook = xlrd.open_workbook(file_contents=data)
    sheets = []
    for sheet in workbook.sheets():
        rows = []
        for row_index in range(min(sheet.nrows, 2000)):
            values = [_stringify_xlrd_cell(sheet.cell(row_index, col_index)) for col_index in range(sheet.ncols)]
            if any(value.strip() for value in values):
                rows.append("\t".join(values))
        if sheet.nrows > 2000:
            rows.append("[...sheet truncated after 2000 rows...]")
        if rows:
            sheets.append(f"--- Sheet: {sheet.name} ---\n" + "\n".join(rows))
    return "\n\n".join(sheets) or "(No data found in legacy spreadsheet)"


def _stringify_xlrd_cell(cell: Any) -> str:
    value = getattr(cell, "value", "")
    ctype = int(getattr(cell, "ctype", 0) or 0)
    if ctype == 0 or value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _run_converter_stdout(command: List[str], data: bytes, ext: str, filename: str) -> str:
    with tempfile.TemporaryDirectory(prefix="metis-parse-") as temp_dir:
        input_path = Path(temp_dir) / _safe_temp_filename(filename, ext)
        input_path.write_bytes(data)
        result = subprocess.run(
            [*command, str(input_path)],
            cwd=temp_dir,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=30,
            check=False,
        )
        if result.returncode != 0:
            raise ValueError(_decode_text(result.stderr) or f"converter exited {result.returncode}")
        return _decode_text(result.stdout).strip() or "(No text extracted by converter)"


def _convert_with_soffice(data: bytes, ext: str, filename: str) -> str:
    soffice = soffice_path()
    if not soffice:
        return ""
    with tempfile.TemporaryDirectory(prefix="metis-parse-") as temp_dir:
        input_path = Path(temp_dir) / _safe_temp_filename(filename, ext)
        input_path.write_bytes(data)
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                temp_dir,
                str(input_path),
            ],
            cwd=temp_dir,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=45,
            check=False,
        )
        txt_candidates = sorted(Path(temp_dir).glob("*.txt"))
        if txt_candidates:
            return _decode_text(txt_candidates[0].read_bytes()).strip() or "(No text extracted by LibreOffice)"
        if result.returncode != 0:
            raise ValueError(_decode_text(result.stderr) or f"LibreOffice exited {result.returncode}")
    return ""


def _safe_temp_filename(filename: str, ext: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", Path(filename or "upload").stem).strip("._") or "upload"
    suffix = ext if ext.startswith(".") else f".{ext}"
    return f"{stem[:48]}{suffix}"


def _dedupe_nonempty(values: Iterable[str]) -> List[str]:
    result: List[str] = []
    previous = ""
    for value in values:
        item = " ".join(str(value or "").split())
        if item and item != previous:
            result.append(item)
            previous = item
    return result
